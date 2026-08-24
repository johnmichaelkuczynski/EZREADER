import os
import logging
import time
import threading
import re
import json
import uuid
from datetime import datetime
from urllib.parse import quote_plus
from flask import Flask, render_template, request, jsonify, session, send_file, send_from_directory, url_for, make_response, Response
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from werkzeug.utils import secure_filename
import requests
import PyPDF2
from docx import Document
from ai_processor import process_text as legacy_process_text, chat_with_ai, split_text
from multi_provider_processor import multi_provider_processor
from PIL import Image
import pytesseract
import io
import base64
from io import BytesIO
from reportlab.pdfgen import canvas
from ai_detector import detect_ai_content # Added import
from humanizer import get_user_profile, add_writing_sample, clear_user_profile, get_user_style_text
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
from style_rewrite_passthrough import process_style_rewrite
from anthropic import Anthropic
import openai

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def preprocess_dollar_signs(text: str) -> str:
    """
    SIMPLE APPROACH: Remove ALL dollar signs from text. Period.
    """
    if not text:
        return text
    
    # SIMPLE: Just remove all dollar signs
    return text.replace('$', '')

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png', 'mp3', 'wav'}
CHUNK_SIZE = 500  # Average words per chunk, approximately 1-2 pages

# Languages for translation
SUPPORTED_LANGUAGES = {
    'en': 'English',
    'es': 'Spanish (Español)',
    'fr': 'French (Français)',
    'de': 'German (Deutsch)',
    'it': 'Italian (Italiano)',
    'pt': 'Portuguese (Português)',
    'nl': 'Dutch (Nederlands)',
    'ru': 'Russian (Русский)',
    'zh': 'Chinese (中文)',
    'ja': 'Japanese (日本語)',
    'ko': 'Korean (한국어)',
    'ar': 'Arabic (العربية)',
    'hi': 'Hindi (हिन्दी)',
    'bn': 'Bengali (বাংলা)',
    'ur': 'Urdu (اردو)',
    'tr': 'Turkish (Türkçe)',
    'vi': 'Vietnamese (Tiếng Việt)',
    'th': 'Thai (ไทย)',
    'id': 'Indonesian (Bahasa Indonesia)',
    'pl': 'Polish (Polski)',
    'cs': 'Czech (Čeština)',
    'hu': 'Hungarian (Magyar)',
    'sv': 'Swedish (Svenska)',
    'fi': 'Finnish (Suomi)',
    'da': 'Danish (Dansk)',
    'no': 'Norwegian (Norsk)',
    'ro': 'Romanian (Română)',
    'uk': 'Ukrainian (Українська)',
    'he': 'Hebrew (עברית)',
    'fa': 'Persian (فارسی)',
    'el': 'Greek (Ελληνικά)'
}

# Languages supported by DeepL API
DEEPL_SUPPORTED_LANGUAGES = [
    'bg', 'cs', 'da', 'de', 'el', 'en', 'es', 'et', 'fi', 'fr', 'hu', 
    'id', 'it', 'ja', 'ko', 'lt', 'lv', 'nb', 'nl', 'pl', 'pt', 'ro', 
    'ru', 'sk', 'sl', 'sv', 'tr', 'uk', 'zh'
]

# Document length modes (in words)
SHORT_DOCUMENT_THRESHOLD = 2000
LONG_DOCUMENT_MAX = 500000

# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 300 * 1024 * 1024

# Configure database connection using environment variables directly
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("NEON_DATABASE_URL") or os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
    "pool_size": 20,  # Increased pool size for better handling of concurrent connections
    "max_overflow": 10  # Allow up to 10 connections beyond pool_size
}

# Configure permanent sessions
app.config['PERMANENT_SESSION_LIFETIME'] = 86400  # 24 hours
app.config['SESSION_TYPE'] = 'filesystem'

db.init_app(app)

with app.app_context():
    try:
        import models
        db.create_all()
        logger.info("Successfully connected to database and created tables")
    except Exception as e:
        # Multiple gunicorn workers may race on CREATE TABLE against a fresh DB;
        # one wins, the others see a UniqueViolation on pg_type. The schema is
        # in place either way, so swallow that specific race and continue.
        from sqlalchemy.exc import IntegrityError, ProgrammingError
        if isinstance(e, (IntegrityError, ProgrammingError)) and "already exists" in str(e).lower():
            logger.warning(f"Database init race (benign — schema exists): {str(e).splitlines()[0]}")
        else:
            logger.error(f"Error initializing database: {str(e)}")
            raise

# Initialize Cross-Chunk Coherence (CC) reconstruction schema in Neon Postgres
try:
    from reconstruction_engine import init_reconstruction_schema, run_reconstruction_stream
    init_reconstruction_schema()
except Exception as e:
    logger.error(f"Reconstruction schema init failed: {e}")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def chunk_text(text):
    """
    Two-level chunking strategy:
    - For documents over 5,000 words: Split into macrochunks of ~5,000 words, each further split into subchunks of ~500 words
    - For smaller documents: Split directly into chunks of ~500 words
    """
    try:
        logger.debug(f"Starting text chunking, total length: {len(text)} characters")
        total_words = len(text.split())
        logger.debug(f"Total word count: {total_words}")
        
        # Determine chunking strategy based on document size
        MACROCHUNK_SIZE = 5000  # words (about 20-25 pages)
        
        # For documents over 5,000 words, use two-level chunking
        if total_words > MACROCHUNK_SIZE:
            logger.debug(f"Document exceeds {MACROCHUNK_SIZE} words, using two-level chunking strategy")
            return two_level_chunking(text, MACROCHUNK_SIZE, CHUNK_SIZE)
        else:
            logger.debug(f"Document is under {MACROCHUNK_SIZE} words, using single-level chunking")
            return single_level_chunking(text, CHUNK_SIZE)
        
    except Exception as e:
        logger.error(f"Error in chunking text: {str(e)}")
        # Fallback to simple chunking in case of error
        simple_chunks = [text[i:i+8000] for i in range(0, len(text), 8000)]
        logger.debug(f"Fallback: Created {len(simple_chunks)} simple chunks")
        return simple_chunks

def single_level_chunking(text, chunk_size):
    """
    Split text into manageable chunks of the specified size.
    
    Ensures that chunks start with complete sentences and maintains
    paragraph integrity where possible.
    """
    try:
        # Split text by paragraphs
        paragraphs = []
        # Handle different paragraph break styles
        for para_candidate in re.split(r'\n\s*\n', text):
            # Further split very long paragraphs
            if len(para_candidate.split()) > chunk_size * 1.5:  # Reduced threshold for better handling
                logger.debug(f"Breaking down very long paragraph: {len(para_candidate.split())} words")
                # Split by sentences for very long paragraphs
                sentences = re.split(r'(?<=[.!?])\s+', para_candidate)
                current_para = []
                current_length = 0
                
                for sentence in sentences:
                    sentence_words = len(sentence.split())
                    
                    # Special handling for extremely long sentences (edge case)
                    if sentence_words > chunk_size and not current_para:
                        # This one sentence exceeds our chunk size and would be the first sentence in the chunk
                        # We need to preserve it but mark it specifically
                        logger.warning(f"Found extremely long sentence: {sentence_words} words")
                        paragraphs.append(f"[LONG_SENTENCE]{sentence}")
                        continue
                        
                    if current_length + sentence_words > chunk_size and current_para:
                        paragraphs.append(' '.join(current_para))
                        current_para = [sentence]
                        current_length = sentence_words
                    else:
                        current_para.append(sentence)
                        current_length += sentence_words
                
                if current_para:
                    paragraphs.append(' '.join(current_para))
            else:
                if para_candidate.strip():
                    paragraphs.append(para_candidate.strip())
        
        logger.debug(f"Split text into {len(paragraphs)} paragraphs")
        
        # Now group paragraphs into chunks, being careful with sentence boundaries
        chunks = []
        current_chunk = []
        current_size = 0

        for paragraph in paragraphs:
            # Special handling for previously marked long sentences
            if paragraph.startswith("[LONG_SENTENCE]"):
                # If we have a current chunk in progress, finish it first
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                
                # Add this long sentence as its own chunk, removing the marker
                chunks.append(paragraph[15:])  # Remove the [LONG_SENTENCE] marker
                continue
                
            words = len(paragraph.split())
            
            # Create a new chunk if adding this paragraph would exceed the limit
            if current_size + words > chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [paragraph]
                current_size = words
            else:
                current_chunk.append(paragraph)
                current_size += words

        # Add the last chunk if there's anything left
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
            
        logger.debug(f"Created {len(chunks)} chunks from text")

        # If no chunks were created, treat the entire text as one chunk
        result = chunks if chunks else [text]
        
        # Log some statistics
        if result:
            chunk_sizes = [len(chunk.split()) for chunk in result]
            logger.debug(f"Chunk sizes (words): min={min(chunk_sizes)}, max={max(chunk_sizes)}, avg={sum(chunk_sizes)/len(chunk_sizes):.1f}")
            
        return result
    except Exception as e:
        logger.error(f"Error in single-level chunking: {str(e)}")
        raise

def two_level_chunking(text, macrochunk_size, subchunk_size):
    """
    Split text into two levels:
    1. First into macrochunks of ~5,000 words
    2. Then each macrochunk into subchunks of ~500 words
    
    Ensures that macrochunks start with complete sentences.
    """
    try:
        # Split text into paragraphs
        paragraphs = []
        for para_candidate in re.split(r'\n\s*\n', text):
            if para_candidate.strip():
                paragraphs.append(para_candidate.strip())
        
        logger.debug(f"Split text into {len(paragraphs)} paragraphs for macrochunking")
        
        # Step 1: Group paragraphs into macrochunks
        macrochunks = []
        current_macrochunk = []
        current_size = 0

        for paragraph in paragraphs:
            words = len(paragraph.split())
            
            # For very long paragraphs, split them into sentences to ensure better chunking
            if words > macrochunk_size and not current_macrochunk:
                # This paragraph alone exceeds the macrochunk size and would be the first item
                # Split it into sentences to ensure complete sentence chunking
                sentences = re.split(r'(?<=[.!?])\s+', paragraph)
                current_sentence_group = []
                sentence_group_size = 0
                
                for sentence in sentences:
                    sentence_words = len(sentence.split())
                    if sentence_group_size + sentence_words > macrochunk_size and current_sentence_group:
                        # Complete the current macrochunk with sentences collected so far
                        macrochunks.append(' '.join(current_sentence_group))
                        current_sentence_group = [sentence]
                        sentence_group_size = sentence_words
                    else:
                        current_sentence_group.append(sentence)
                        sentence_group_size += sentence_words
                
                # Add any remaining sentences as their own macrochunk
                if current_sentence_group:
                    macrochunks.append(' '.join(current_sentence_group))
            
            # Normal case: Check if adding this paragraph would exceed the macrochunk size
            elif current_size + words > macrochunk_size and current_macrochunk:
                macrochunks.append('\n\n'.join(current_macrochunk))
                current_macrochunk = [paragraph]
                current_size = words
            else:
                current_macrochunk.append(paragraph)
                current_size += words

        # Add the last macrochunk if there's anything left
        if current_macrochunk:
            macrochunks.append('\n\n'.join(current_macrochunk))
            
        logger.debug(f"Created {len(macrochunks)} macrochunks from text")
        
        # Step 2: Split each macrochunk into subchunks
        all_subchunks = []
        for i, macrochunk in enumerate(macrochunks, 1):
            logger.debug(f"Processing macrochunk {i} of {len(macrochunks)}")
            subchunks = single_level_chunking(macrochunk, subchunk_size)
            all_subchunks.extend(subchunks)
        
        logger.debug(f"Created {len(all_subchunks)} total subchunks from {len(macrochunks)} macrochunks")
        
        # Log some statistics on the subchunks
        if all_subchunks:
            chunk_sizes = [len(chunk.split()) for chunk in all_subchunks]
            logger.debug(f"Subchunk sizes (words): min={min(chunk_sizes)}, max={max(chunk_sizes)}, avg={sum(chunk_sizes)/len(chunk_sizes):.1f}")
        
        return all_subchunks if all_subchunks else [text]
    except Exception as e:
        logger.error(f"Error in two-level chunking: {str(e)}")
        raise

def extract_text_from_pdf(file_path):
    """Extract text from PDF with better error handling and Unicode cleaning"""
    import re
    text = []
    try:
        logger.debug(f"Starting PDF extraction from {file_path}")
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean Unicode surrogate characters and invalid characters
                        cleaned_text = re.sub(r'[\uD800-\uDFFF]', '', page_text)
                        cleaned_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', cleaned_text)
                        # Ensure UTF-8 compatibility
                        cleaned_text = cleaned_text.encode('utf-8', errors='ignore').decode('utf-8')
                        text.append(cleaned_text)
                    else:
                        logger.warning(f"No text extracted from page {page_num}")
                except Exception as page_error:
                    logger.error(f"Error extracting text from page {page_num}: {str(page_error)}")
                    continue

        if not text:
            raise ValueError("No text could be extracted from the PDF")

        # Apply dollar sign preprocessing at document ingestion
        extracted_text = '\n\n'.join(text)
        return preprocess_dollar_signs(extracted_text)
    except Exception as e:
        logger.error(f"Error in PDF extraction: {str(e)}")
        raise ValueError(f"Could not extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from DOCX with improved handling for complex documents"""
    try:
        logger.debug(f"Starting DOCX extraction from {file_path}")
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        # Handle headers and footers if available
        try:
            for section in doc.sections:
                # Try to access header and footer
                try:
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                full_text.append(f"Header: {para.text}")
                except Exception as header_err:
                    logger.debug(f"Could not extract header: {str(header_err)}")
                
                try:
                    if section.footer:
                        for para in section.footer.paragraphs:
                            if para.text.strip():
                                full_text.append(f"Footer: {para.text}")
                except Exception as footer_err:
                    logger.debug(f"Could not extract footer: {str(footer_err)}")
        except Exception as section_err:
            logger.debug(f"Error processing sections: {str(section_err)}")
        
        # Join all text with double line breaks
        result = '\n\n'.join(full_text)
        logger.debug(f"Extracted {len(full_text)} text elements from DOCX")
        
        if not result.strip():
            raise ValueError("No text could be extracted from the document")
            
        # Apply dollar sign preprocessing at document ingestion
        return preprocess_dollar_signs(result)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")

def extract_text_from_image(image_path):
    """Extract text from an image using OCR"""
    try:
        image = Image.open(image_path)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        text = pytesseract.image_to_string(image)
        # Apply dollar sign preprocessing at document ingestion
        return preprocess_dollar_signs(text.strip())
    except Exception as e:
        logger.error(f"Error extracting text from image: {str(e)}")
        raise

def extract_text_from_audio(audio_path):
    """Extract text from audio file using OpenAI's Whisper API"""
    try:
        logger.debug(f"Starting audio transcription from {audio_path}")
        
        # Check file exists
        if not os.path.exists(audio_path):
            raise ValueError(f"Audio file not found: {audio_path}")
        
        # Import the whisper transcription module
        from whisper_transcription import transcribe_audio_with_whisper
        
        # Use Whisper API for transcription
        text = transcribe_audio_with_whisper(audio_path)
        
        if not text or text.strip() == "":
            raise ValueError("Whisper API returned empty transcription")
            
        # Apply dollar sign preprocessing at document ingestion
        return preprocess_dollar_signs(text)
    except ImportError:
        logger.error("Failed to import whisper_transcription module")
        raise ValueError("Audio transcription service is not properly configured")
    except Exception as e:
        logger.error(f"Error in Whisper audio transcription: {str(e)}")
        raise ValueError(f"Could not extract text from audio: {str(e)}")

def process_chunk(chunk_id, chunk_number, custom_instructions='', is_first_chunk=False, email=None, author_style='', content_source=None, ai_provider=None, preserve_length=True):
    """
    Process a single chunk with given instructions and optional user style
    using the multi-provider processor for improved reliability and performance
    
    Args:
        chunk_id: ID of the document containing the chunk
        chunk_number: Number of the chunk to process
        custom_instructions: Additional processing instructions
        is_first_chunk: Whether this is the first chunk (for style handling)
        email: User email for personal style
        author_style: Specific author style to imitate (e.g., "Orwell")
        content_source: Optional content source text to incorporate
        ai_provider: Optional AI provider to use (openai, anthropic, perplexity)
        preserve_length: Whether to maintain original text length (default: True)
    """
    start_time = time.time()
    
    # Get the chunk from the database
    chunk = models.DocumentChunk.query.filter_by(
        document_id=chunk_id,
        chunk_number=chunk_number
    ).first()

    if not chunk:
        raise ValueError('Chunk not found')
        
    # Update status to processing
    chunk.processing_status = 'processing'
    db.session.commit()

    try:
        # Get the document to check if it has a user profile
        document = models.TextEntry.query.get(chunk_id)
        
        # Get user style if available
        user_style_text = None
        if document and document.user_profile_id:
            user_style_text = get_user_style_text(profile_id=document.user_profile_id)
        elif email:
            user_style_text = get_user_style_text(email=email)
        
        # Get content source if available
        content_source_text = ""
        content_source_instructions = ""
        
        # First, check if a content source was provided directly via API
        if content_source:
            logger.debug(f"Using content source provided via API ({len(content_source)} characters)")
            content_source_text = content_source
        else:
            # If not, check for content sources in the database
            content_sources = models.ContentSource.query.filter_by(text_entry_id=chunk_id).all()
            if content_sources:
                logger.debug(f"Found {len(content_sources)} content sources for document {chunk_id}")
                for source in content_sources:
                    content_source_text += source.text_content + "\n\n"
                    if source.usage_instructions:
                        content_source_instructions += source.usage_instructions + "\n\n"
                
                logger.debug(f"Content source length from database: {len(content_source_text)}")
                if content_source_instructions:
                    logger.debug(f"Content source instructions: {content_source_instructions}")
        
        # If we have content source text from any source, log it
        if content_source_text:
            logger.debug(f"Total content source length: {len(content_source_text)} characters")
            
        # Check if this is the first chunk or a subsequent chunk
        include_style_in_output = is_first_chunk

        # Use the new multi-provider processor for better reliability
        logger.debug(f"Processing chunk {chunk_number} with multi-provider processor")
        
        # Generate effective instructions
        effective_instructions = custom_instructions
        
        # Add content source instructions if available
        if content_source_text:
            if not content_source_instructions:
                content_source_instructions = "Intelligently enrich the target document with relevant information, ideas, examples, and arguments from the content source, without overriding the target document's structure or identity."
            
            # Add content source instructions to custom instructions
            additional_instructions = (
                f"\n\nCONTENT SOURCE INSTRUCTIONS: {content_source_instructions}\n\n"
                f"CONTENT SOURCE TEXT:\n{content_source_text}\n\n"
                "Use the CONTENT SOURCE TEXT to enrich the Target Document (the input text) according to the CONTENT SOURCE INSTRUCTIONS."
            )
            
            if effective_instructions:
                effective_instructions += additional_instructions
            else:
                effective_instructions = additional_instructions
        
        # Add author style if specified
        if author_style:
            if effective_instructions:
                effective_instructions += f". Write in the style of {author_style}"
            else:
                effective_instructions = f"Write in the style of {author_style}"
            logger.debug(f"Added author style '{author_style}' to instructions")
        
        # Add mandatory length preservation instructions
        length_preservation_instructions = """
MANDATORY LENGTH PRESERVATION: Your rewritten output MUST match or exceed the length of the original text.
Target is 100-110% of the original word count.
Under NO circumstances should your output be shorter than the input.
"""
        if effective_instructions:
            effective_instructions = length_preservation_instructions + "\n" + effective_instructions
        else:
            effective_instructions = length_preservation_instructions
            
        try:
            # First, get the original word count
            original_word_count = len(chunk.original_chunk.split())
            logger.info(f"Original chunk word count: {original_word_count}")
            
            # Check if length multiplier is specified in custom instructions
            length_multiplier = None
            if '3x' in custom_instructions.lower() or 'triple' in custom_instructions.lower():
                length_multiplier = 3.0
            elif '2x' in custom_instructions.lower() or 'double' in custom_instructions.lower():
                length_multiplier = 2.0
            elif '1.5x' in custom_instructions.lower() or '150%' in custom_instructions.lower():
                length_multiplier = 1.5
                
            if length_multiplier:
                logger.info(f"Detected length multiplier: {length_multiplier}x in custom instructions")
            
            # Process the text - passing content_source directly to the processor rather than through instructions
            # Also pass the AI provider preference and maintain_length parameter
            provider_preference = None
            if ai_provider and ai_provider in ['openai', 'anthropic', 'perplexity']:
                provider_preference = ai_provider
                logger.info(f"Using user-selected AI provider: {provider_preference}")

            # Log content source usage for debugging
            if content_source_text:
                logger.info(f"Using content source with {len(content_source_text)} characters")
            
            # Apply dollar sign preprocessing to chunk before processing
            processed_chunk_text = preprocess_dollar_signs(chunk.original_chunk)
            
            processed_text = multi_provider_processor.process_text(
                text=processed_chunk_text,
                action='rewrite',
                custom_instructions=effective_instructions,
                include_style_in_output=include_style_in_output,
                user_style_text=user_style_text,
                content_source=content_source_text if content_source_text else None,
                length_multiplier=length_multiplier,
                provider_preference=provider_preference,
                maintain_length=preserve_length
            )
            
            # Also apply post-processing to ensure no dollar signs escape
            processed_text = preprocess_dollar_signs(processed_text)
            
            # Verify length requirements
            processed_word_count = len(processed_text.split())
            length_ratio = processed_word_count / original_word_count
            logger.info(f"Processed chunk word count: {processed_word_count}, ratio: {length_ratio:.2f}")
            
            # If output is shorter than input, force expansion
            max_expansion_attempts = 3
            current_attempt = 0
            
            while length_ratio < 1.0 and current_attempt < max_expansion_attempts:
                current_attempt += 1
                logger.warning(f"Output too short ({length_ratio:.2f}). Emergency expansion attempt {current_attempt}/{max_expansion_attempts}")
                
                emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {processed_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications
4. DO NOT add unrelated material or irrelevant content
5. DO NOT summarize - this is a rewrite with expansion

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                
                # Try to expand the text (maintaining content_source for consistency)
                processed_text = multi_provider_processor.process_text(
                    text=processed_text,
                    action='rewrite',
                    custom_instructions=emergency_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text,
                    content_source=content_source_text if content_source_text else None,
                    length_multiplier=length_multiplier
                )
                
                # Re-check length requirements
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                logger.info(f"After expansion attempt {current_attempt}: {processed_word_count} words, ratio: {length_ratio:.2f}")
                
                if length_ratio >= 1.0:
                    logger.info(f"Length requirement met after {current_attempt} expansion attempts")
                    break
            
            # If still too short after all attempts, make one final aggressive attempt
            if length_ratio < 1.0:
                logger.warning(f"Still below minimum length ({length_ratio:.2f}) after {max_expansion_attempts} attempts. Final attempt.")
                
                final_emergency_instructions = f"""
CRITICAL LENGTH FAILURE - FINAL ATTEMPT:
Your output is MUCH TOO SHORT. The original text had {original_word_count} words, but your rewrite has only {processed_word_count} words.

You MUST expand this text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - your primary task is to meet the length requirement.

For EACH paragraph:
1. DOUBLE the size with detailed examples, evidence, and elaboration
2. Add substantial explanations for each concept
3. Explore the implications and applications of each idea
4. Maintain all intellectual arguments and flow from the original

DO NOT summarize. DO NOT condense. EXPAND with meaningful content.
"""
                
                processed_text = multi_provider_processor.process_text(
                    text=processed_text,
                    action='rewrite',
                    custom_instructions=final_emergency_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text,
                    content_source=content_source_text if content_source_text else None,
                    length_multiplier=length_multiplier
                )
                
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                logger.info(f"After final expansion attempt: {processed_word_count} words, ratio: {length_ratio:.2f}")
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
            # Fallback to legacy processor if the new processor fails
            processed_text = legacy_process_text(
                text=chunk.original_chunk,
                action='rewrite',
                custom_instructions=effective_instructions if 'effective_instructions' in locals() else custom_instructions,
                include_style_in_output=include_style_in_output,
                user_style_text=user_style_text
            )
            
            # Even with legacy processor, verify length requirements
            original_word_count = len(chunk.original_chunk.split())
            processed_word_count = len(processed_text.split())
            length_ratio = processed_word_count / original_word_count
            
            if length_ratio < 1.0:
                logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion instructions.")
                emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving all meaning and intellectual depth."
                
                try:
                    processed_text = legacy_process_text(
                        text=processed_text,
                        action='expand',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False,
                        user_style_text=user_style_text
                    )
                except Exception as expand_error:
                    logger.error(f"Error expanding with legacy processor: {str(expand_error)}")
                    # Continue with the best we have at this point

        # Check for timeout
        if time.time() - start_time > 45:  # 45-second timeout
            raise TimeoutError(f"Processing chunk {chunk_number} timed out")
            
        # Translate if target language is specified
        if document and document.target_language and document.target_language != 'en':
            logger.debug(f"Translating to {document.target_language}")
            try:
                # Track the original length for validation
                original_word_count = len(processed_text.split())
                logger.debug(f"Original text word count before translation: {original_word_count}")
                
                # Use our enhanced chunking translation
                translated_text = translate_text(processed_text, document.target_language)
                processed_text = translated_text
                
                # Validate the translation was complete
                translated_word_count = len(processed_text.split())
                logger.debug(f"Translation completed, length: {len(processed_text)}, word count: {translated_word_count}")
                
                # Verify we didn't lose a significant amount of content
                if translated_word_count < original_word_count * 0.6 and translated_word_count < 100:
                    logger.error(f"Translation validation failed - output too short: {translated_word_count}/{original_word_count} words")
                    # Send alert to logging but continue with what we have
            except Exception as translation_error:
                logger.error(f"Translation error: {str(translation_error)}")
                # Continue with untranslated text if translation fails

        chunk.processed_chunk = processed_text
        chunk.is_processed = True
        chunk.processing_status = 'complete'
        db.session.commit()
        
        logger.debug(f"Successfully processed chunk {chunk_number} in {time.time() - start_time:.2f} seconds")
        
        return chunk
        
    except Exception as e:
        # Get detailed error information
        import traceback
        error_traceback = traceback.format_exc()
        
        # Mark the chunk as having an error
        chunk.processing_status = 'error'
        chunk.processed_chunk = f"[Error processing chunk: {str(e)}]"
        db.session.commit()
        
        # Log both the error message and full traceback
        logger.error(f"Error processing chunk {chunk_number}: {str(e)}")
        logger.error(f"Full traceback for chunk {chunk_number}: {error_traceback}")
        
        # Re-raise the exception to be handled by the calling function
        raise

@app.route('/')
def index():
    return render_template('index_clean.html')
    
@app.route('/translate')
def translation_page():
    """Render the dedicated translation page with source and target language boxes"""
    return render_template('translation_page.html')

@app.route('/clean')
def clean_design():
    """Render the new clean design starting with header and upload"""
    return render_template('index_clean.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not supported'}), 400

        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                text = extract_text_from_docx(file_path)
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                text = extract_text_from_image(file_path)
            elif filename.lower().endswith(('.mp3', '.wav')):
                text = extract_text_from_audio(file_path)
            else:
                return jsonify({'error': 'Unsupported file format'}), 400

            # Split text into chunks
            chunks = chunk_text(text)
            if not chunks:
                return jsonify({'error': 'No text could be extracted'}), 400

            logger.debug(f"Created {len(chunks)} chunks from the text")

            # Create text entry and chunks
            text_entry = models.TextEntry(
                original_text=text,
                processed_text="",  # Will be built from chunks
                action="rewrite",
                complexity="default",
                total_chunks=len(chunks)
            )
            db.session.add(text_entry)
            db.session.flush()

            # Create chunks
            for i, chunk in enumerate(chunks, 1):
                doc_chunk = models.DocumentChunk(
                    document_id=text_entry.id,
                    chunk_number=i,
                    original_chunk=chunk,
                    processed_chunk="",
                    processing_status = "pending"
                )
                db.session.add(doc_chunk)

            db.session.commit()

            # Store document ID in session
            session['current_document_id'] = text_entry.id
            session['current_chunk'] = 1
            
            # Calculate document stats
            word_count = len(text.split())
            chars_count = len(text)
            ave_chunk_size = word_count // len(chunks) if len(chunks) > 0 else 0

            # Send the ENTIRE text, not just the first chunk
            return jsonify({
                'text': text,  # Send the complete original text
                'document_id': text_entry.id,
                'total_chunks': len(chunks),
                'current_chunk': 1,
                'file_name': filename,
                'total_words': word_count,
                'total_chars': chars_count,
                'average_chunk_size': ave_chunk_size,
                'full_document_loaded': True  # Flag to indicate the full document was processed
            })

        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)

    except Exception as e:
        logger.error(f"Error in file upload: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/get_chunk', methods=['POST', 'GET'])
def get_chunk():
    try:
        # Support both GET and POST methods
        if request.method == 'GET':
            document_id = request.args.get('document_id')
            chunk_number = request.args.get('chunk_number', 1, type=int)
            get_all = request.args.get('all', 'false').lower() == 'true'
        else:
            data = request.get_json()
            document_id = data.get('document_id')
            chunk_number = data.get('chunk_number', 1)
            get_all = data.get('all', False)

        if not document_id:
            return jsonify({'error': 'No document selected'}), 400
            
        # Get the document
        document = models.TextEntry.query.get(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404
            
        # If 'all' parameter is true, return all chunks
        if get_all:
            logger.debug(f"Retrieving all chunks for document {document_id}")
            chunks = models.DocumentChunk.query.filter_by(
                document_id=document_id
            ).order_by(models.DocumentChunk.chunk_number).all()
            
            if not chunks:
                return jsonify({'error': 'No chunks found for document'}), 404
                
            return jsonify({
                'chunks': [chunk.to_dict() for chunk in chunks],
                'total_chunks': document.total_chunks
            })
        else:
            # Return a single chunk
            chunk = models.DocumentChunk.query.filter_by(
                document_id=document_id,
                chunk_number=chunk_number
            ).first()

            if not chunk:
                return jsonify({'error': 'Chunk not found'}), 404

            return jsonify({
                'chunk': chunk.to_dict(),
                'total_chunks': document.total_chunks
            })

    except Exception as e:
        logger.error(f"Error getting chunk: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/process_chunk', methods=['POST'])
def process_chunk_route():
    try:
        data = request.get_json()
        document_id = data.get('document_id')
        chunk_number = data.get('chunk_number')
        custom_instructions = data.get('custom_instructions', '')
        is_first_chunk = data.get('is_first_chunk', False)
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        target_language = data.get('target_language', '')
        content_source = data.get('content_source', '')  # Get content source from the request
        ai_provider = data.get('ai_provider', '')  # Get user's AI provider preference
        preserve_length = data.get('preserve_length', True)  # Get length preservation setting

        if not document_id or not chunk_number:
            return jsonify({'error': 'Missing document_id or chunk_number'}), 400

        # If no email provided, try to get from session
        if not email:
            email = session.get('last_email')
            
        # Get the document to save target language if provided
        if target_language:
            document = models.TextEntry.query.get(document_id)
            if document:
                document.target_language = target_language
                db.session.commit()
                logger.debug(f"Set target language for document {document_id} to {target_language}")
        
        # If content_source is provided in the request, add it to the instructions
        content_source_instructions = ''
        if content_source:
            # Log the length to avoid logging the entire content source which could be large
            logger.debug(f"Received content source in request: {len(content_source)} characters")
            
            # Add a specific marker in the instructions for content source if not already included
            if 'content source' not in custom_instructions.lower() and 'use the provided content' not in custom_instructions.lower():
                content_source_instructions = " Use the provided content source to enrich the text."
            
            # Add the content source instruction to custom_instructions
            custom_instructions = custom_instructions.strip() + content_source_instructions

        try:
            # Pass all parameters to the process_chunk function including content_source, ai_provider, and preserve_length
            chunk = process_chunk(
                document_id, 
                chunk_number, 
                custom_instructions, 
                is_first_chunk, 
                email,
                author_style,  # Add author style parameter
                content_source,  # Pass content source to process_chunk
                ai_provider,  # Pass the selected AI provider
                preserve_length  # Pass length preservation setting
            )
            return jsonify({'chunk': chunk.to_dict()})
        except ValueError as e:
            return jsonify({'error': str(e)}), 404

    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Error in process_chunk_route: {str(e)}")
        logger.error(f"Traceback in process_chunk_route: {error_traceback}")
        return jsonify({'error': str(e), 'traceback': error_traceback}), 500

@app.route('/process_all_chunks', methods=['POST'])
def process_all_chunks():
    try:
        data = request.get_json()
        document_id = data.get('document_id')
        custom_instructions = data.get('custom_instructions', '')
        # Add support for author style
        author_style = data.get('author_style', '')
        email = data.get('email', '')
        target_language = data.get('target_language', '')

        if not document_id:
            return jsonify({'error': 'Missing document_id'}), 400
            
        # Get the document to save target language if provided
        if target_language:
            document = models.TextEntry.query.get(document_id)
            if document:
                document.target_language = target_language
                db.session.commit()
                logger.debug(f"Set target language for document {document_id} to {target_language}")

        # Get all chunks for the document
        chunks = models.DocumentChunk.query.filter_by(
            document_id=document_id
        ).order_by(models.DocumentChunk.chunk_number).all()

        if not chunks:
            return jsonify({'error': 'No chunks found'}), 404

        total_chunks = len(chunks)
        processed_chunks = 0
        errors = []

        # Find the next unprocessed chunk
        current_chunk = None
        for chunk in chunks:
            if not chunk.is_processed:
                current_chunk = chunk
                break

        if not current_chunk:
            # All chunks are processed
            return jsonify({
                'status': 'complete',
                'processed_chunks': total_chunks,
                'total_chunks': total_chunks,
                'percentage': 100
            })

        try:
            # Process the current chunk with timeout
            start_time = time.time()

            # Mark chunk as processing
            current_chunk.processing_status = "processing"
            db.session.commit()

            # Determine if this is the first chunk - only chunk #1 should include style instructions
            is_first_chunk = current_chunk.chunk_number == 1
            
            logger.debug(f"Processing chunk {current_chunk.chunk_number}, is_first_chunk={is_first_chunk}")

            # Get the document to check if it has a user profile
            document = models.TextEntry.query.get(document_id)
            
            # Get user style if available
            user_style_text = None
            if document and document.user_profile_id:
                user_style_text = get_user_style_text(profile_id=document.user_profile_id)
            
            # Use email from session if available
            email = session.get('last_email')
            if not user_style_text and email:
                user_style_text = get_user_style_text(email=email)

            # Use the new multi-provider processor for better reliability
            try:
                # If author style is specified, include it in custom instructions with maximum emphasis
                effective_instructions = custom_instructions
                if author_style:
                    # CRITICAL: This must be the first instruction and given highest priority
                    author_style_instruction = f"!!! ABSOLUTELY MANDATORY !!! STRICTLY WRITE IN THE EXACT STYLE OF {author_style.upper()} - THIS IS THE HIGHEST PRIORITY INSTRUCTION AND MUST BE FOLLOWED EXACTLY !!!"
                    
                    if effective_instructions:
                        effective_instructions = author_style_instruction + "\n\n" + effective_instructions
                    else:
                        effective_instructions = author_style_instruction
                # Pass user instructions without automatic modifications
                # No automatic length or conciseness instructions that override user requirements
                    
                # First, get the original word count
                original_word_count = len(current_chunk.original_chunk.split())
                logger.info(f"Original chunk word count: {original_word_count}")
                
                # Determine if user selected a specific AI provider
                provider_preference = None
                if ai_provider and ai_provider in ['openai', 'anthropic', 'perplexity']:
                    provider_preference = ai_provider
                    logger.info(f"Using user-selected AI provider: {provider_preference}")
                
                processed_text = multi_provider_processor.process_text(
                    text=current_chunk.original_chunk,
                    action='rewrite',
                    custom_instructions=effective_instructions,
                    include_style_in_output=is_first_chunk,
                    user_style_text=user_style_text,
                    provider_preference=provider_preference,
                    maintain_length=preserve_length
                )
                
                # Verify length requirements
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                logger.info(f"Processed chunk word count: {processed_word_count}, ratio: {length_ratio:.2f}")
                
                # If output is shorter than input, force expansion
                if length_ratio < 1.0:
                    logger.warning(f"Output too short ({length_ratio:.2f}). Performing emergency expansion.")
                    
                    emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {processed_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                    
                    # Try to expand the text
                    processed_text = multi_provider_processor.process_text(
                        text=processed_text,
                        action='rewrite',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False,
                        user_style_text=user_style_text
                    )
                    
                    # Re-check length
                    processed_word_count = len(processed_text.split())
                    length_ratio = processed_word_count / original_word_count
                    logger.info(f"After expansion: {processed_word_count} words, ratio: {length_ratio:.2f}")
                
            except Exception as e:
                logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
                # Fallback to legacy processor if the new processor fails
                processed_text = legacy_process_text(
                    text=current_chunk.original_chunk,
                    action='rewrite',
                    custom_instructions=effective_instructions if 'effective_instructions' in locals() else custom_instructions,
                    include_style_in_output=is_first_chunk,
                    user_style_text=user_style_text
                )
                
                # Even with legacy processor, verify length requirements
                if 'original_word_count' not in locals():
                    original_word_count = len(current_chunk.original_chunk.split())
                
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                
                if length_ratio < 1.0:
                    logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion.")
                    try:
                        emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving meaning."
                        
                        processed_text = legacy_process_text(
                            text=processed_text,
                            action='expand',
                            custom_instructions=emergency_instructions,
                            include_style_in_output=False
                        )
                    except Exception as expand_error:
                        logger.error(f"Error expanding with legacy processor: {str(expand_error)}")

            # Check for timeout
            if time.time() - start_time > 30:  # Increased timeout to 30 seconds
                raise TimeoutError(f"Processing chunk {current_chunk.chunk_number} timed out")

            if "Error processing text:" in processed_text:
                raise Exception(processed_text)

            # Update chunk status
            current_chunk.processed_chunk = processed_text
            current_chunk.is_processed = True
            current_chunk.processing_status = "complete"
            db.session.commit()

            # Count actually processed chunks
            processed_chunks = models.DocumentChunk.query.filter_by(
                document_id=document_id,
                is_processed=True
            ).count()

            # Return progress
            percentage = int((processed_chunks / total_chunks) * 100)
            return jsonify({
                'status': 'in_progress',
                'processed_chunks': processed_chunks,
                'total_chunks': total_chunks,
                'percentage': percentage,
                'current_chunk': current_chunk.chunk_number,
                'processed_text': processed_text,
                'errors': errors if errors else None
            })

        except Exception as chunk_error:
            error_msg = str(chunk_error)
            logger.error(f"Error processing chunk {current_chunk.chunk_number}: {error_msg}")

            # Mark chunk for retry
            current_chunk.processing_status = "error"
            current_chunk.is_processed = False  # Reset for retry
            db.session.commit()

            errors.append(f"Chunk {current_chunk.chunk_number}: {error_msg}")

            # Return error status
            return jsonify({
                'error': error_msg,
                'processed_chunks': processed_chunks,
                'total_chunks': total_chunks,
                'percentage': int((processed_chunks / total_chunks) * 100),
                'current_chunk': current_chunk.chunk_number,
                'status': 'error',
                'retry': True
            }), 500

    except Exception as e:
        logger.error(f"Error in process_all_chunks: {str(e)}")
        return jsonify({
            'error': str(e),
            'processed_chunks': processed_chunks if 'processed_chunks' in locals() else 0,
            'total_chunks': total_chunks if 'total_chunks' in locals() else 0,
            'percentage': int((processed_chunks / total_chunks) * 100) if 'processed_chunks' in locals() and 'total_chunks' in locals() and total_chunks > 0 else 0
        }), 500

@app.route('/process', methods=['POST'])
def process():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        text = data.get('text', '')
        custom_instructions = data.get('custom_instructions', '')
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        ai_provider = data.get('ai_provider', '')
        preserve_length = data.get('preserve_length', True)
        target_language = data.get('target_language', '')
        content_source = data.get('content_source', '')
        processing_mode = data.get('processing_mode', 'rewrite')

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        # CRITICAL: Apply dollar sign preprocessing before any AI interaction
        text = preprocess_dollar_signs(text)
        logger.debug(f"Applied dollar sign preprocessing to input text")
        
        # Also preprocess content source if provided
        if content_source:
            content_source = preprocess_dollar_signs(content_source)
            logger.debug(f"Applied dollar sign preprocessing to content source")

        # Initialize user_style_text for all modes
        user_style_text = None
        if email:
            user_style_text = get_user_style_text(email=email)
        else:
            # Try to get from session
            email = session.get('last_email')
            if email:
                user_style_text = get_user_style_text(email=email)

        # Handle different processing modes
        if processing_mode == 'homework':
            # Homework mode: treat input text as instructions to follow
            instructions_to_follow = text
            
            # Build homework completion prompt
            homework_prompt = f"""You are an AI assistant helping to complete a task, assignment, or follow instructions. The following text contains the instructions or questions you need to address:

INSTRUCTIONS/TASK:
{instructions_to_follow}

Please complete this task thoroughly and professionally. Follow all instructions exactly as given. If it's an exam, answer all questions completely. If it's homework, solve all problems with explanations. If it's a set of instructions, follow them precisely.

Additional guidance: {custom_instructions if custom_instructions else 'None provided.'}
"""
            
            # For homework mode, we use the prompt as the text to process
            text = homework_prompt
            
            # Clear custom instructions since they're now part of the main prompt
            custom_instructions = ""
            
            # Homework mode doesn't need length preservation or author style
            preserve_length = False
            author_style = ""
            
            logger.info(f"Processing in homework mode with {len(instructions_to_follow)} characters of instructions")
        else:
            # Rewrite mode: traditional text transformation
            # Add author style to custom instructions if provided, with MAXIMUM emphasis
            if author_style:
                # CRITICAL: This must be the first instruction and given highest priority
                author_style_instruction = f"!!! ABSOLUTELY MANDATORY !!! STRICTLY WRITE IN THE EXACT STYLE OF {author_style.upper()} - THIS IS THE HIGHEST PRIORITY INSTRUCTION AND MUST BE FOLLOWED EXACTLY !!!"
                
                if custom_instructions:
                    custom_instructions = author_style_instruction + "\n\n" + custom_instructions
                else:
                    custom_instructions = author_style_instruction
                logger.debug(f"Added author style '{author_style}' to instructions with maximum emphasis")

            # Add mandatory length preservation instructions for rewrite mode
            if preserve_length:
                length_preservation_instructions = """
MANDATORY LENGTH PRESERVATION: Your rewritten output MUST match or exceed the length of the original text.
Target is 100-110% of the original word count.
Under NO circumstances should your output be shorter than the input.
"""
                if custom_instructions:
                    custom_instructions = length_preservation_instructions + "\n" + custom_instructions
                else:
                    custom_instructions = length_preservation_instructions
            
        # Process the text with the new multi-provider processor
        try:
            # First, get the original word count
            original_word_count = len(text.split())
            logger.info(f"Original text word count: {original_word_count}")
            
            # Add content source instructions if available
            if content_source:
                logger.debug(f"Content source provided for direct text processing: {len(content_source)} characters")
                
                # Add content source instructions to custom instructions
                content_source_instructions = "Intelligently enrich the target document with relevant information, ideas, examples, and arguments from the content source, without overriding the target document's structure or identity."
                
                additional_instructions = (
                    f"\n\nCONTENT SOURCE INSTRUCTIONS: {content_source_instructions}\n\n"
                    f"CONTENT SOURCE TEXT:\n{content_source}\n\n"
                    "Use the CONTENT SOURCE TEXT to enrich the Target Document (the input text) according to the CONTENT SOURCE INSTRUCTIONS."
                )
                
                if custom_instructions:
                    custom_instructions += additional_instructions
                else:
                    custom_instructions = additional_instructions
                
                logger.debug("Added content source instructions to direct processing")
            
            # Determine if user selected a specific AI provider
            provider_preference = None
            if ai_provider and ai_provider in ['openai', 'anthropic', 'perplexity', 'azure']:
                provider_preference = ai_provider
                logger.info(f"Using user-selected AI provider: {provider_preference}")
            
            result = multi_provider_processor.process_text(
                text=text,
                action='rewrite',
                custom_instructions=custom_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text,
                maintain_length=preserve_length,  # Use the user's length preservation preference
                provider_preference=provider_preference  # Use the user's AI provider preference
            )
            
            # Verify length requirements
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            logger.info(f"Processed text word count: {result_word_count}, ratio: {length_ratio:.2f}")
            
            # If output is shorter than input, force expansion
            if length_ratio < 1.0:
                logger.warning(f"Output too short ({length_ratio:.2f}). Performing emergency expansion.")
                
                emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {result_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                
                # Try to expand the text
                result = multi_provider_processor.process_text(
                    text=result,
                    action='rewrite',
                    custom_instructions=emergency_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text
                )
                
                # Re-check length
                result_word_count = len(result.split())
                length_ratio = result_word_count / original_word_count
                logger.info(f"After expansion: {result_word_count} words, ratio: {length_ratio:.2f}")
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
            # Fallback to legacy processor if the new processor fails
            result = legacy_process_text(
                text=text,
                action='rewrite',
                custom_instructions=custom_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text
            )
            
            # Even with legacy processor, verify length requirements
            if 'original_word_count' not in locals():
                original_word_count = len(text.split())
            
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            
            if length_ratio < 1.0:
                logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion.")
                try:
                    emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving meaning."
                    
                    result = legacy_process_text(
                        text=result,
                        action='expand',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False
                    )
                except Exception as expand_error:
                    logger.error(f"Error expanding with legacy processor: {str(expand_error)}")
            
        # Translate the result if target language is specified
        if target_language and target_language != 'en':
            logger.debug(f"Translating result to {target_language}")
            try:
                # Track original word count for validation
                original_word_count = len(result.split())
                logger.debug(f"Original word count before translation: {original_word_count}")
                
                # Use our chunked translation system
                translated_result, engine_used = translate_text(result, target_language)
                result = translated_result
                
                # Validate the translation output
                translated_word_count = len(result.split())
                logger.debug(f"Translation completed, length: {len(result)}, word count: {translated_word_count}")
                
                # Check if we lost significant content (accounting for language differences)
                if translated_word_count < original_word_count * 0.6:
                    logger.warning(f"Translation appears shorter than expected: {translated_word_count}/{original_word_count} words")
                    if translated_word_count < 100 and original_word_count > 500:
                        logger.error("Translation may have been truncated. Output too short compared to input.")
            except Exception as translation_error:
                logger.error(f"Translation error: {str(translation_error)}")
                # Continue with untranslated text if translation fails

        return jsonify({'result': result})
    except Exception as e:
        logger.error(f"Error in /process: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/combine_target_source', methods=['POST'])
def combine_target_source():
    """Directly combine target and source text with instructions without saving"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        target_text = data.get('target_text', '')
        source_text = data.get('source_text', '')
        source_instructions = data.get('source_instructions', '')
        custom_instructions = data.get('custom_instructions', '')
        target_language = data.get('target_language', '')
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        
        if not target_text:
            return jsonify({'error': 'No target text provided'}), 400
            
        if not source_text:
            return jsonify({'error': 'No source text provided'}), 400
        
        # Set default source instructions if not provided
        if not source_instructions:
            source_instructions = "Intelligently enrich the target document with relevant information, ideas, examples, and arguments from the content source, without overriding the target document's structure or identity."
        
        # Get user style if available
        user_style_text = None
        if email:
            user_style_text = get_user_style_text(email=email)
        
        # Add author style to custom instructions if provided
        if author_style:
            if custom_instructions:
                custom_instructions += f". Write in the style of {author_style}"
            else:
                custom_instructions = f"Write in the style of {author_style}"
        
        # Combine instructions for processing
        effective_instructions = custom_instructions
        
        # Add content source instructions
        additional_instructions = (
            f"\n\nCONTENT SOURCE INSTRUCTIONS: {source_instructions}\n\n"
            f"CONTENT SOURCE TEXT:\n{source_text}\n\n"
            "Use the CONTENT SOURCE TEXT to enrich the Target Document (the input text) according to the CONTENT SOURCE INSTRUCTIONS."
        )
        
        if effective_instructions:
            effective_instructions += additional_instructions
        else:
            effective_instructions = additional_instructions
            
        # Add mandatory length preservation instructions
        length_preservation_instructions = """
MANDATORY LENGTH PRESERVATION: Your rewritten output MUST match or exceed the length of the original text.
Target is 100-110% of the original word count.
Under NO circumstances should your output be shorter than the input.
"""
        if effective_instructions:
            effective_instructions = length_preservation_instructions + "\n" + effective_instructions
        else:
            effective_instructions = length_preservation_instructions
            
        # Process the text with combined instructions
        try:
            # First, get the original word count
            original_word_count = len(target_text.split())
            logger.info(f"Original target text word count: {original_word_count}")
            
            result = multi_provider_processor.process_text(
                text=target_text,
                action='rewrite',
                custom_instructions=effective_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text
            )
            
            # Verify length requirements
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            logger.info(f"Processed text word count: {result_word_count}, ratio: {length_ratio:.2f}")
            
            # If output is shorter than input, force expansion
            if length_ratio < 1.0:
                logger.warning(f"Output too short ({length_ratio:.2f}). Performing emergency expansion.")
                
                emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {result_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications
4. Using MORE content from the source text to enrich the output

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                
                # Try to expand the text
                result = multi_provider_processor.process_text(
                    text=result,
                    action='rewrite',
                    custom_instructions=emergency_instructions + "\n\n" + additional_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text
                )
                
                # Re-check length
                result_word_count = len(result.split())
                length_ratio = result_word_count / original_word_count
                logger.info(f"After expansion: {result_word_count} words, ratio: {length_ratio:.2f}")
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
            # Fallback to legacy processor if the new processor fails
            result = legacy_process_text(
                text=target_text,
                action='rewrite',
                custom_instructions=effective_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text
            )
            
            # Even with legacy processor, verify length requirements
            if 'original_word_count' not in locals():
                original_word_count = len(target_text.split())
            
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            
            if length_ratio < 1.0:
                logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion.")
                try:
                    emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving meaning and using relevant content from the source text for enrichment."
                    
                    result = legacy_process_text(
                        text=result,
                        action='expand',
                        custom_instructions=emergency_instructions + "\n\n" + additional_instructions,
                        include_style_in_output=False
                    )
                except Exception as expand_error:
                    logger.error(f"Error expanding with legacy processor: {str(expand_error)}")
        
        # Translate if a target language is specified
        if target_language and target_language != 'en':
            logger.debug(f"Translating result to {target_language}")
            try:
                # Track original word count for validation
                original_word_count = len(result.split())
                logger.debug(f"Original word count before translation: {original_word_count}")
                
                # Use our chunked translation system
                translated_result, engine_used = translate_text(result, target_language)
                result = translated_result
                
                # Validate the translation output
                translated_word_count = len(result.split())
                logger.debug(f"Translation completed, length: {len(result)}, word count: {translated_word_count}")
                
                # Check if we lost significant content (accounting for language differences)
                if translated_word_count < original_word_count * 0.6:
                    logger.warning(f"Translation appears shorter than expected: {translated_word_count}/{original_word_count} words")
                    if translated_word_count < 100 and original_word_count > 500:
                        logger.error("Translation may have been truncated. Output too short compared to input.")
            except Exception as translation_error:
                logger.error(f"Translation error: {str(translation_error)}")
                # Continue with untranslated text if translation fails
            
        return jsonify({'result': result})
    except Exception as e:
        logger.error(f"Error combining target and source: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        message = data.get('message', '')
        context = data.get('context', '')

        if not message:
            return jsonify({'error': 'No message provided'}), 400

        logger.debug("Processing chat message")
        response = chat_with_ai(message, context)

        # Save to database
        chat_entry = models.ChatMessage(
            message=message,
            response=response,
            context=context
        )
        db.session.add(chat_entry)
        db.session.commit()

        return jsonify({'response': response})
    except Exception as e:
        logger.error(f"Error in /chat: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/history', methods=['GET'])
def get_history():
    try:
        text_entries = models.TextEntry.query.order_by(models.TextEntry.created_at.desc()).limit(10).all()
        chat_messages = models.ChatMessage.query.order_by(models.ChatMessage.created_at.desc()).limit(10).all()

        history = {
            'text_entries': [entry.to_dict() for entry in text_entries],
            'chat_messages': [msg.to_dict() for msg in chat_messages]
        }
        return jsonify(history)
    except Exception as e:
        logger.error(f"Error fetching history: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/extract_text', methods=['POST'])
def extract_text():
    try:
        # Check if this is a file upload (FormData) or JSON
        if request.files and 'file' in request.files:
            # Handle file upload
            file = request.files['file']
            if file and file.filename:
                filename = secure_filename(file.filename)
                file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
                
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(temp_path)
                
                try:
                    # Extract text based on file type
                    if file_ext == 'pdf':
                        extracted_text = extract_text_from_pdf(temp_path)
                    elif file_ext in ['docx', 'doc']:
                        extracted_text = extract_text_from_docx(temp_path)
                    elif file_ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']:
                        extracted_text = extract_text_from_image(temp_path)
                    elif file_ext == 'txt':
                        with open(temp_path, 'r', encoding='utf-8') as f:
                            extracted_text = preprocess_dollar_signs(f.read())
                    else:
                        return jsonify({'success': False, 'error': f'Unsupported file type: {file_ext}'}), 400
                    
                    return jsonify({'success': True, 'text': extracted_text})
                finally:
                    # Clean up temp file
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
            else:
                return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        # Handle JSON requests (legacy)
        data = request.get_json()
        image_data = data.get('image_data', '')
        text_data = data.get('text', '')

        # If image data is provided, extract text from image
        if image_data:
            # Save image data to a temporary file
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_image.png')
            try:
                image_bytes = base64.b64decode(image_data)
                image = Image.open(io.BytesIO(image_bytes))
                image.save(temp_path)
                extracted_text = extract_text_from_image(temp_path)
                return jsonify({'text': extracted_text})
            except Exception as e:
                logger.error(f"Error extracting text from image: {str(e)}")
                return jsonify({'error': 'Error extracting text from image'}), 500
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        # If text data is provided, create a document from the text
        elif text_data:
            try:
                # Split the text into chunks
                chunks = chunk_text(text_data)
                
                # Create text entry and chunks
                text_entry = models.TextEntry(
                    original_text=text_data,
                    processed_text="",  # Will be built from chunks
                    action="rewrite",
                    complexity="default",
                    total_chunks=len(chunks)
                )
                db.session.add(text_entry)
                db.session.flush()

                # Create chunks
                for i, chunk in enumerate(chunks, 1):
                    doc_chunk = models.DocumentChunk(
                        document_id=text_entry.id,
                        chunk_number=i,
                        original_chunk=chunk,
                        processed_chunk="",
                        processing_status="pending"
                    )
                    db.session.add(doc_chunk)

                db.session.commit()
                
                # Return document information
                return jsonify({
                    'document_id': text_entry.id,
                    'total_chunks': len(chunks),
                    'current_chunk': 1
                })
                
            except Exception as e:
                logger.error(f"Error creating document from text: {str(e)}")
                db.session.rollback()
                return jsonify({'error': 'Error creating document from text'}), 500
        else:
            return jsonify({'error': 'No text or image data provided'}), 400

    except Exception as e:
        logger.error(f"Error in text extraction: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/download_document/<format>', methods=['POST'])
def download_document(format):
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        filename = f'processed_text.{format}'
        
        if format == 'pdf':
            mime_type = 'application/pdf'
        elif format == 'docx':
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif format == 'latex':
            mime_type = 'text/plain'
            filename = 'processed_text.tex'
        else:
            return jsonify({'error': 'Unsupported format'}), 400

        buffer = BytesIO()

        if format == 'pdf':
            # Create PDF with proper text wrapping
            pdf = canvas.Canvas(buffer)
            width = 500  # Maximum width in points (leaving margins)
            y = 800  # Start from top of page
            font_name = 'Helvetica'
            font_size = 12
            pdf.setFont(font_name, font_size)

            # Process each paragraph
            for paragraph in text.split('\n'):
                if not paragraph.strip():
                    y -= 15  # Add space between paragraphs
                    continue

                words = paragraph.split()
                line = []

                for word in words:
                    line.append(word)
                    line_width = pdf.stringWidth(' '.join(line), font_name, font_size)

                    if line_width > width:
                        # Remove last word as it caused overflow
                        line.pop()
                        # Draw the line
                        pdf.drawString(50, y, ' '.join(line))
                        y -= 15
                        # Start new line with the overflow word
                        line = [word]

                    # Check if we need a new page
                    if y < 50:
                        pdf.showPage()
                        pdf.setFont(font_name, font_size)
                        y = 800

                # Draw remaining text in the line
                if line:
                    pdf.drawString(50, y, ' '.join(line))
                    y -= 15

            pdf.save()

        elif format == 'docx':
            # Create Word document
            doc = Document()
            doc.add_paragraph(text)
            doc.save(buffer)
        elif format == 'latex':
            # Create LaTeX document
            latex_content = f"""\\documentclass{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{amsmath}}
\\usepackage{{amsfonts}}
\\usepackage{{amssymb}}
\\usepackage{{geometry}}
\\geometry{{margin=1in}}

\\title{{Processed Document}}
\\author{{EZ Reader}}
\\date{{\\today}}

\\begin{{document}}
\\maketitle

{text}

\\end{{document}}"""
            buffer.write(latex_content.encode('utf-8'))
        else:
            return jsonify({'error': 'Unsupported format'}), 400

        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype=mime_type
        )

    except Exception as e:
        logger.error(f"Error in download_{format}: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/detect_ai', methods=['POST']) # Added AI detection route
def detect_ai():
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        logger.info("AI detection request received, text length: %d", len(text))

        if not text:
            logger.warning("Empty text submitted for AI detection")
            return jsonify({'error': 'No text provided'}), 400
            
        # Add some random variance to the text to ensure different results
        # This helps avoid caching on GPTZero's side
        import random
        import string
        # Add a hidden character at the end with timestamp
        hidden_suffix = f"\n<!-- {int(time.time())}-{random.randint(1000, 9999)} -->"
        modified_text = text + hidden_suffix
        
        logger.info("Sending text to AI detector with unique suffix")
        result = detect_ai_content(modified_text)

        if 'error' in result:
            # Check if it's a configuration error (missing API key)
            if result.get('is_configuration_error', False):
                logger.error("Configuration error in AI detection: %s", result['error'])
                # Return a 400 error for configuration issues
                return jsonify({
                    'error': result['error'],
                    'is_configuration_error': True
                }), 400
            else:
                logger.error("Service error in AI detection: %s", result['error'])
                # Return a 500 error for service failures
                return jsonify({'error': result['error']}), 500
        
        # Add additional user-friendly information
        if 'document_class' in result:
            if result['document_class'] == 'ai':
                result['conclusion'] = 'This text appears to be AI-generated.'
            elif result['document_class'] == 'human':
                result['conclusion'] = 'This text appears to be human-written.'
            elif result['document_class'] == 'mixed':
                result['conclusion'] = 'This text appears to contain a mixture of AI and human-written content.'
            else:
                result['conclusion'] = 'Unable to determine if this text is AI-generated or human-written.'
        
        # Remove raw response to reduce payload size
        if 'raw_response' in result:
            del result['raw_response']
            
        logger.info("AI detection complete, score: %s, class: %s", 
                   result.get('ai_score', 'unknown'), 
                   result.get('document_class', 'unknown'))

        # Add 'score' field for frontend compatibility
        if 'ai_score' in result:
            result['score'] = result['ai_score']

        return jsonify(result)

    except Exception as e:
        logger.error(f"Error in AI detection endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.errorhandler(500)
def handle_500_error(error):
    return jsonify({'error': 'Internal server error occurred. Please try again later.'}), 500

@app.errorhandler(404)
def handle_404_error(error):
    return jsonify({'error': 'Requested resource not found'}), 404

@app.route('/get_last_email', methods=['GET'])
def get_last_email():
    """Get the last used email from the session for auto-fill"""
    last_email = session.get('last_email', '')
    return jsonify({'email': last_email})

@app.route('/translate', methods=['POST'])
def translate_text():
    """Simple translation endpoint that handles both small and large documents"""
    try:
        # Get translation parameters from request
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        text = data.get('text', '')
        source_language = data.get('source_language', 'auto')
        target_language = data.get('target_language', 'en')
        ai_provider = data.get('ai_provider', 'openai')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
            
        # Use our simple translation service
        from simple_translation import translate_text as perform_translation
        
        # Perform the translation
        translated_text, metadata = perform_translation(
            text=text,
            source_language=source_language,
            target_language=target_language,
            ai_provider=ai_provider
        )
        
        # Check for errors
        if 'error' in metadata:
            return jsonify({'error': metadata['error']}), 500
            
        # Return the translation result with metadata
        return jsonify({
            'result': translated_text,
            'engine_used': metadata.get('engine_used', ai_provider),
            'elapsed_seconds': metadata.get('elapsed_seconds', 0),
            'words_per_second': metadata.get('words_per_second', 0),
            'word_count': metadata.get('word_count', 0),
            'message': metadata.get('message', '')
        })
            
    except Exception as e:
        logger.error(f"Error in translate_text: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/share_rewrite', methods=['POST'])
def share_rewrite():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        email = data.get('email')
        text = data.get('text')
        subject = data.get('subject', 'Your Rewritten Text')
        
        if not email or not text:
            return jsonify({'error': 'Email and text are required'}), 400
            
        # Validate email format
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            return jsonify({'error': 'Invalid email format'}), 400
        
        # Store the user's email in session for future use
        session['last_email'] = email
        
        logger.info(f"Sending email to {email} with subject: {subject}")
        
        # Import email service
        from email_service import send_text_email
        
        # Send email with appropriate format based on content length
        success, message = send_text_email(
            to_email=email,
            subject=subject,
            text=text
        )
        
        if success:
            return jsonify({'success': True, 'message': 'Email sent successfully'})
        else:
            logger.error(f"Failed to send email: {message}")
            return jsonify({'error': message}), 500
        
    except Exception as e:
        logger.error(f"Error sharing rewrite: {str(e)}")
        return jsonify({'error': f'Error sharing rewrite: {str(e)}'}), 500

@app.route('/api/humanizer/profile', methods=['POST'])
def create_profile():
    """Create or get a user profile by email"""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
            
        try:
            profile = get_user_profile(email)
            return jsonify({
                'success': True,
                'profile': profile.to_dict()
            })
        except Exception as e:
            logger.error(f"Error creating/getting user profile: {str(e)}")
            return jsonify({'error': f'Profile error: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"Error in create_profile: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/humanizer/upload', methods=['POST'])
def upload_writing_sample():
    """Upload a writing sample and add to user profile"""
    try:
        # Check if file is included
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
            
        # Get profile_id from form data
        profile_id = request.form.get('profile_id')
        if not profile_id:
            email = request.form.get('email')
            if not email:
                return jsonify({'error': 'Either profile_id or email is required'}), 400
            profile = get_user_profile(email)
            profile_id = profile.id
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
            
        # Process the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                text = extract_text_from_docx(file_path)
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                text = extract_text_from_image(file_path)
            elif filename.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return jsonify({'error': 'Unsupported file format'}), 400
                
            # Determine file type
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'txt'
            
            # Add writing sample
            sample = add_writing_sample(
                profile_id=profile_id,
                filename=filename,
                text_content=text,
                file_type=ext
            )
            
            return jsonify({
                'success': True,
                'sample': sample.to_dict()
            })
            
        except Exception as e:
            logger.error(f"Error processing writing sample: {str(e)}")
            return jsonify({'error': str(e)}), 500
        finally:
            # Clean up the file
            if os.path.exists(file_path):
                os.remove(file_path)
                
    except Exception as e:
        logger.error(f"Error in upload_writing_sample: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/humanizer/samples', methods=['GET'])
def get_writing_samples():
    """Get all writing samples for a user profile"""
    try:
        profile_id = request.args.get('profile_id')
        email = request.args.get('email')
        
        if not profile_id and not email:
            return jsonify({'error': 'Either profile_id or email is required'}), 400
            
        if email and not profile_id:
            profile = get_user_profile(email)
            profile_id = profile.id
            
        # Get samples
        samples = models.WritingSample.query.filter_by(profile_id=profile_id).all()
        
        return jsonify({
            'success': True,
            'samples': [sample.to_dict() for sample in samples],
            'count': len(samples)
        })
        
    except Exception as e:
        logger.error(f"Error getting writing samples: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/upload', methods=['POST'])
def upload_content_source():
    """Upload a content source document for text enrichment"""
    try:
        # Log debugging information
        logger.debug(f"Content source upload request received: {request.files}")
        logger.debug(f"Form data: {request.form}")
        
        # Check if file is included
        if 'file' not in request.files:
            logger.error("No file found in request.files")
            return jsonify({'error': 'No file provided'}), 400
            
        # Get text_entry_id from form data - now optional
        text_entry_id = request.form.get('text_entry_id')
        usage_instructions = request.form.get('usage_instructions', '')
        
        logger.debug(f"text_entry_id: {text_entry_id}, instructions length: {len(usage_instructions)}")
        
        # If text_entry_id is provided, verify it exists
        text_entry = None
        if text_entry_id:
            text_entry = models.TextEntry.query.get(text_entry_id)
            if not text_entry and text_entry_id:
                logger.error(f"Text entry not found for ID: {text_entry_id}")
                return jsonify({'error': 'Text entry not found'}), 404
            
        file = request.files['file']
        if file.filename == '':
            logger.error("Empty filename received")
            return jsonify({'error': 'No file selected'}), 400
        
        logger.debug(f"File received: {file.filename}")
            
        # Process the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        logger.debug(f"File saved to: {file_path}")
        
        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                text = extract_text_from_docx(file_path)
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                text = extract_text_from_image(file_path)
            elif filename.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return jsonify({'error': 'Unsupported file format'}), 400
                
            # Determine file type
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'txt'
            
            # Calculate word count
            word_count = len(text.split())
            
            # Add content source
            content_source = models.ContentSource(
                text_entry_id=text_entry_id,
                filename=filename,
                text_content=text,
                word_count=word_count,
                file_type=ext,
                usage_instructions=usage_instructions,
                created_at=datetime.utcnow()
            )
            
            db.session.add(content_source)
            db.session.commit()
            
            # Log success
            logger.debug(f"Successfully created content source: {content_source.id}, text length: {len(text)}")
            
            return jsonify({
                'success': True,
                'content_source': {
                    'id': content_source.id,
                    'filename': content_source.filename,
                    'word_count': content_source.word_count,
                    'file_type': content_source.file_type,
                    'text_entry_id': content_source.text_entry_id,
                    'text_content': text
                }
            })
            
        except Exception as e:
            logger.error(f"Error processing content source file: {str(e)}")
            return jsonify({'error': str(e)}), 500
        finally:
            # Clean up the file
            if os.path.exists(file_path):
                os.remove(file_path)
                
    except Exception as e:
        logger.error(f"Error in upload_content_source: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/save_instructions', methods=['POST'])
def save_content_source_instructions():
    """Save usage instructions for a content source"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        content_source_id = data.get('content_source_id')
        usage_instructions = data.get('usage_instructions', '')
        
        if not content_source_id:
            return jsonify({'error': 'content_source_id is required'}), 400
            
        # Find the content source
        content_source = models.ContentSource.query.get(content_source_id)
        if not content_source:
            return jsonify({'error': 'Content source not found'}), 404
            
        # Update instructions
        content_source.usage_instructions = usage_instructions
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Content source instructions updated successfully'
        })
        
    except Exception as e:
        logger.error(f"Error in save_content_source_instructions: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/get', methods=['GET'])
def get_content_sources():
    """Get all content sources for a text entry"""
    try:
        text_entry_id = request.args.get('text_entry_id')
        
        if not text_entry_id:
            return jsonify({'error': 'text_entry_id is required'}), 400
        
        logger.debug(f"Getting content sources for text_entry_id: {text_entry_id}")
            
        # Get content sources
        content_sources = models.ContentSource.query.filter_by(text_entry_id=text_entry_id).all()
        
        logger.debug(f"Found {len(content_sources)} content sources")
        
        return jsonify({
            'success': True,
            'content_sources': [{
                'id': source.id,
                'filename': source.filename,
                'word_count': source.word_count,
                'file_type': source.file_type,
                'usage_instructions': source.usage_instructions,
                'created_at': source.created_at.isoformat(),
                'text_entry_id': source.text_entry_id
            } for source in content_sources],
            'count': len(content_sources)
        })
    except Exception as e:
        logger.error(f"Error getting content sources: {str(e)}")
        return jsonify({'error': str(e)}), 500
        
@app.route('/api/content_source/get_text', methods=['GET'])
def get_content_source_text():
    """Get the text content of a specific content source"""
    try:
        content_source_id = request.args.get('content_source_id')
        
        if not content_source_id:
            return jsonify({'error': 'content_source_id is required'}), 400
            
        logger.debug(f"Getting text content for content_source_id: {content_source_id}")
            
        # Get content source
        content_source = models.ContentSource.query.get(content_source_id)
        
        if not content_source:
            logger.error(f"Content source not found for ID: {content_source_id}")
            return jsonify({'error': 'Content source not found'}), 404
            
        logger.debug(f"Found content source: {content_source.filename} with {len(content_source.text_content)} chars")
            
        return jsonify({
            'success': True,
            'text_content': content_source.text_content,
            'filename': content_source.filename,
            'file_type': content_source.file_type
        })
        
    except Exception as e:
        logger.error(f"Error getting content sources: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/save_text', methods=['POST'])
def save_content_source_text():
    """Save text content directly as a content source"""
    try:
        # Log debugging information
        logger.debug(f"Content source text save request received: {request.json}")
        
        data = request.json
        if not data or 'text_content' not in data:
            logger.error("No text content provided in request")
            return jsonify({'error': 'No text content provided'}), 400
        
        text_content = data.get('text_content')
        if not text_content or not text_content.strip():
            logger.error("Empty text content provided")
            return jsonify({'error': 'Text content cannot be empty'}), 400
            
        filename = data.get('filename', 'pasted_content.txt')
        text_entry_id = data.get('text_entry_id')
        usage_instructions = data.get('usage_instructions', '')
        
        logger.debug(f"text_entry_id: {text_entry_id}, instructions length: {len(usage_instructions)}")
        
        # If text_entry_id is provided, verify it exists
        if text_entry_id:
            text_entry = models.TextEntry.query.get(text_entry_id)
            if not text_entry:
                logger.error(f"Text entry not found for ID: {text_entry_id}")
                return jsonify({'error': 'Text entry not found'}), 404
        
        # Calculate word count
        word_count = len(text_content.split())
        
        # Create content source
        content_source = models.ContentSource(
            text_entry_id=text_entry_id,
            filename=filename,
            text_content=text_content,
            word_count=word_count,
            file_type='txt',
            usage_instructions=usage_instructions,
            created_at=datetime.utcnow()
        )
        
        db.session.add(content_source)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'content_source': {
                'id': content_source.id,
                'filename': content_source.filename,
                'word_count': content_source.word_count,
                'file_type': content_source.file_type,
                'text_entry_id': content_source.text_entry_id
            }
        })
        
    except Exception as e:
        logger.error(f"Error in save_content_source_text: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/delete', methods=['POST'])
def delete_content_source():
    """Delete a content source"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        content_source_id = data.get('content_source_id')
        
        if not content_source_id:
            return jsonify({'error': 'content_source_id is required'}), 400
            
        # Find the content source
        content_source = models.ContentSource.query.get(content_source_id)
        if not content_source:
            return jsonify({'error': 'Content source not found'}), 404
            
        # Delete the content source
        db.session.delete(content_source)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Content source deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error in delete_content_source: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/humanizer/clear', methods=['POST'])
def clear_profile():
    """Clear all writing samples for a user profile"""
    try:
        data = request.get_json()
        profile_id = data.get('profile_id')
        email = data.get('email')
        
        if not profile_id and not email:
            return jsonify({'error': 'Either profile_id or email is required'}), 400
            
        if email and not profile_id:
            profile = get_user_profile(email)
            profile_id = profile.id
            
        # Clear profile
        profile = clear_user_profile(profile_id)
        
        return jsonify({
            'success': True,
            'message': 'Profile cleared successfully',
            'profile': profile.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error clearing profile: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/process_audio', methods=['POST'])
def process_audio():
    """Process audio file uploads and return transcribed text"""
    file_path = None
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Check if it's an audio file
        file_ext = os.path.splitext(file.filename)[1].lower()
        supported_audio_formats = ['.mp3', '.wav', '.m4a', '.mp4', '.mpeg', '.mpga', '.webm', '.flac', '.aac', '.ogg']
        if file_ext not in supported_audio_formats:
            return jsonify({'error': f'Unsupported audio format. Please use: {", ".join(supported_audio_formats)}'}), 400

        # Create a unique filename to avoid conflicts
        import uuid
        timestamp = int(time.time())
        unique_filename = f"audio_{timestamp}_{uuid.uuid4().hex[:8]}{file_ext}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        
        # Ensure upload directory exists
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Save audio file
        file.save(file_path)
        logger.info(f"Audio file saved to: {file_path}")
        
        # Verify file was saved and has content
        if not os.path.exists(file_path):
            return jsonify({'error': 'Failed to save audio file'}), 500
            
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return jsonify({'error': 'Uploaded audio file is empty'}), 400
            
        logger.info(f"Audio file size: {file_size} bytes")

        try:
            # Extract text from audio
            text = extract_text_from_audio(file_path)
            
            if not text or not text.strip():
                return jsonify({'error': 'No speech could be transcribed from the audio'}), 400
                
            logger.info(f"Audio transcription successful, text length: {len(text)}")
            return jsonify({'text': text.strip()})
            
        except Exception as e:
            logger.error(f"Error processing audio file: {str(e)}")
            return jsonify({'error': f'Error processing audio: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"Error in process_audio: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up the uploaded file
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Cleaned up audio file: {file_path}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up audio file: {cleanup_error}")

@app.route('/get_audio_file/<filename>', methods=['GET'])
def get_audio_file(filename):
    """Serve an audio file for streaming in the browser"""
    if not (filename.startswith('audiobook_') or filename.startswith('tts_')) or not filename.endswith('.mp3'):
        return jsonify({'error': 'Invalid filename format'}), 400
        
    try:
        uploads_dir = app.config['UPLOAD_FOLDER']
        response = send_from_directory(uploads_dir, filename, as_attachment=False, mimetype='audio/mpeg')
        response.headers['Content-Type'] = 'audio/mpeg'
        return response
    except Exception as e:
        logger.error(f"Error serving audio file: {str(e)}")
        return jsonify({'error': f"Could not serve audio file: {str(e)}"}), 404

@app.route('/download_audio_file/<filename>', methods=['GET'])
def download_audio_file(filename):
    """Download an audio file to the user's device"""
    if not (filename.startswith('audiobook_') or filename.startswith('tts_')) or not filename.endswith('.mp3'):
        return jsonify({'error': 'Invalid filename format'}), 400
        
    try:
        uploads_dir = app.config['UPLOAD_FOLDER']
        # For download, use attachment to prompt download with correct MP3 headers
        response = send_from_directory(
            uploads_dir, 
            filename, 
            as_attachment=True,
            download_name=f"Rewritten_Text_Audiobook.mp3",
            mimetype='audio/mpeg'
        )
        response.headers['Content-Type'] = 'audio/mpeg'
        response.headers['Content-Disposition'] = 'attachment; filename="Rewritten_Text_Audiobook.mp3"'
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['Content-Description'] = 'File Transfer'
        response.headers['Cache-Control'] = 'no-cache, must-revalidate'
        return response
    except Exception as e:
        logger.error(f"Error downloading audio file: {str(e)}")
        return jsonify({'error': f"Could not download audio file: {str(e)}"}), 404

@app.route('/download_static_audio/<filename>')
def download_static_audio(filename):
    """Direct download route for audio files from static directory"""
    return send_from_directory('static/audio', filename, as_attachment=True)
        
@app.route('/create_audiobook', methods=['POST'])
def create_audiobook_route():
    """Convert text to audiobook using Azure Speech TTS with language detection"""
    try:
        # Extract request data
        data = request.json
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
            
        text = data['text']
        # Use a shorter text if quota is limited (optional parameter)
        use_reduced_length = data.get('use_reduced_length', False)
        
        # If text is very long and we're asked to reduce length, only use first part
        word_count = len(text.split())
        original_word_count = word_count  # Save for logging/reporting
        
        if use_reduced_length:
            # For reduced length, use approximately first 300 words
            if word_count > 300:
                words = text.split()
                shortened_text = ' '.join(words[:300])
                # Try to find a sentence end for a clean cut
                if '.' in shortened_text:
                    last_period = shortened_text.rindex('.')
                    text = shortened_text[:last_period+1]
                else:
                    text = shortened_text
                
                logger.info(f"Reduced text from {original_word_count} to {len(text.split())} words to avoid processing very large text")
        
        # Remove dollar signs from text before TTS
        text = preprocess_dollar_signs(text)
        
        # Get voice and speed parameters
        voice = data.get('voice', 'nova')
        speed = data.get('speed', 1.0)
        gender = 'female'  # Default gender for compatibility
        
        # Fallback: If old gender parameter is used without voice, map to OpenAI voices
        if 'gender' in data and 'voice' not in data:
            gender = data.get('gender', 'female')
            voice_mapping = {
                'female': 'nova',    # Warm female voice
                'male': 'echo',      # Clear male voice  
                'neutral': 'alloy'   # Neutral voice
            }
            voice = voice_mapping.get(gender, 'nova')
        elif 'gender' in data:
            gender = data.get('gender', 'female')
        
        # Map voice back to gender for compatibility
        if voice in ['nova', 'shimmer']:
            gender = 'female'
        elif voice in ['echo', 'fable', 'onyx']:
            gender = 'male'
        else:
            gender = 'neutral'
        
        # Ensure voice is always set
        if not voice:
            voice = 'nova'  # Default fallback voice
        
        # Import OpenAI TTS
        from openai_tts import convert_text_to_speech
        
        # Create the audiobook with OpenAI TTS
        success, result, file_path = convert_text_to_speech(
            text=text,
            voice=voice,
            speed=speed
        )
        
        if not success:
            # Log the full error for debugging
            logger.error(f"OpenAI TTS audiobook creation failed: {result}")
            
            # Check for quota exceeded error  
            if isinstance(result, str) and any(term in result.lower() for term in ["quota", "credit", "exceed", "limit", "rate"]):
                return jsonify({
                    'error': 'OpenAI TTS quota exceeded. Please try using the "Use Minimal Text" option or use a shorter text.',
                    'error_type': 'quota_exceeded',
                    'word_count': original_word_count
                }), 429  # Use 429 Too Many Requests for quota issues
            
            # Provide user-friendly error message
            return jsonify({
                'error': str(result),
                'error_type': 'synthesis_error'
            }), 500
            
        # Handle new return format with both filenames
        if isinstance(result, dict):
            filename = result['filename']
            static_filename = result['static_filename']
        else:
            # Fallback for old format
            filename = result
            static_filename = result
            
        # Return properly formatted URLs using the actual request host
        # Get the actual host from the request to ensure correct domain
        host = request.headers.get('Host', 'localhost:5000')
        base_url = f"https://{host}"
        audio_url = f"{base_url}/get_audio_file/{filename}"
        download_url = f"{base_url}/get_audio_file/{filename}"
        
        return jsonify({
            'audio_url': audio_url,
            'download_url': download_url,
            'filename': static_filename if isinstance(result, dict) else result,
            'voice': voice,
            'gender': gender,
            'narrator': f"{voice.capitalize()} Voice (OpenAI HD)",
            'message': 'Created using OpenAI TTS HD',
            'model': 'tts-1-hd'
        })
        
    except Exception as e:
        logger.error(f"Error creating audiobook: {str(e)}")
        return jsonify({'error': str(e), 'error_type': 'general_error'}), 500

@app.route('/get_language_voices', methods=['GET'])
def get_language_voices():
    """Get the list of available voices by language from Azure Speech"""
    try:
        from azure_tts import get_language_voices as get_azure_voices
        
        # Get voice data from Azure
        language_voices = get_azure_voices()
        
        # Transform into a format suitable for the frontend
        result = {}
        for lang_code, voices in language_voices.items():
            # Skip languages with no voices
            if not voices:
                continue
                
            # Get the language name if available
            language_name = None
            from azure_tts import LANGUAGE_NAMES
            if lang_code in LANGUAGE_NAMES:
                language_name = LANGUAGE_NAMES[lang_code]
            else:
                language_name = f"Language {lang_code}"
                
            # Add to result
            result[lang_code] = {
                'name': language_name,
                'voices': voices
            }
            
        return jsonify({
            'languages': result,
            'provider': 'azure',
            'status': 'success'
        })
        
    except Exception as e:
        logger.error(f"Error fetching language voices: {str(e)}")
        return jsonify({
            'error': f"Could not fetch voices: {str(e)}",
            'status': 'error'
        }), 500

@app.route('/rewrite_from_output', methods=['POST'])
def rewrite_from_output():
    """
    Process a rewrite request directly from the output text with additional critique instructions.
    This allows users to refine and regenerate text without copying back to the input box.
    
    Request JSON parameters:
    - text: The current output text to be rewritten
    - critique: The user's critique and instructions for the rewrite
    - author_style: Optional author style to emulate
    - content_source: Optional content source to incorporate
    - email: Optional email for personal style
    - ai_provider: Optional AI provider to use (openai, anthropic, perplexity)
    - preserve_length: Optional boolean to determine if length preservation is required
    
    Returns:
    - JSON response with the rewritten text
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        text = data.get('text', '')
        critique = data.get('critique', '')
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        content_source = data.get('content_source', '')
        ai_provider = data.get('ai_provider', '')
        preserve_length = data.get('preserve_length', True)  # Default to preserving length
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
            
        if not critique:
            return jsonify({'error': 'No critique or instructions provided'}), 400
        
        logger.debug(f"Processing rewrite from output with {len(text)} chars of text and {len(critique)} chars of critique")
        logger.debug(f"Using AI provider: {ai_provider if ai_provider else 'default'}")
        logger.debug(f"Preserve length: {preserve_length}")
        
        # Get user style if available
        user_style_text = None
        if email:
            user_style_text = get_user_style_text(email=email)
        else:
            # Try to get from session
            email = session.get('last_email')
            if email:
                user_style_text = get_user_style_text(email=email)
        
        # Create rewrite instructions that emphasize the critique should be applied to existing text
        rewrite_instructions = f"""CRITIQUE REWRITE INSTRUCTIONS:

The text below has been reviewed by the user who wants the following changes:

{critique}

Apply ONLY these requested changes to the text. DO NOT rewrite sections that don't need to be changed.
Focus specifically on addressing the critique while maintaining the overall structure and content of the original text.
"""
        
        # Add author style with MAXIMUM emphasis if provided
        if author_style:
            # CRITICAL: This must come first for highest priority
            author_style_instruction = f"!!! ABSOLUTELY MANDATORY !!! STRICTLY WRITE IN THE EXACT STYLE OF {author_style.upper()} - THIS IS THE HIGHEST PRIORITY INSTRUCTION AND MUST BE FOLLOWED EXACTLY !!!"
            rewrite_instructions = author_style_instruction + "\n\n" + rewrite_instructions
        
        # Add content source instructions if available
        if content_source:
            logger.debug(f"Content source provided for critique rewrite: {len(content_source)} characters")
            
            # Add content source instructions to rewrite instructions
            content_source_instructions = "Intelligently enrich the output with relevant information from the content source while addressing the critique."
            
            additional_instructions = (
                f"\n\nCONTENT SOURCE INSTRUCTIONS: {content_source_instructions}\n\n"
                f"CONTENT SOURCE TEXT:\n{content_source}\n\n"
                "Use the CONTENT SOURCE TEXT to enhance your response according to the critique while maintaining the core content."
            )
            
            rewrite_instructions += additional_instructions
            logger.debug("Added content source instructions to critique rewrite")
        
        # Pass user instructions without automatic modifications
        # User instructions take absolute priority without interference
        
        # Process the text with the multi-provider processor
        try:
            # First, get the original word count
            original_word_count = len(text.split())
            logger.info(f"Original text word count for critique rewrite: {original_word_count}")
            
            # Process with full critique instructions, using specified provider if available
            result = multi_provider_processor.process_text(
                text=text,
                action='rewrite',
                custom_instructions=rewrite_instructions,
                include_style_in_output=False,  # No need for style prefix in output
                user_style_text=user_style_text,
                provider_preference=ai_provider if ai_provider else None
            )
            
            # Only check length requirements if length preservation is requested
            if preserve_length:
                result_word_count = len(result.split())
                length_ratio = result_word_count / original_word_count
                logger.info(f"Critique rewrite word count: {result_word_count}, ratio: {length_ratio:.2f}")
                
                # If output is shorter than input, force expansion
                if length_ratio < 1.0:
                    logger.warning(f"Critique rewrite too short ({length_ratio:.2f}). Performing emergency expansion.")
                    
                    emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {result_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                    
                    # Try to expand the text
                    result = multi_provider_processor.process_text(
                        text=result,
                        action='rewrite',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False,
                        user_style_text=user_style_text,
                        provider_preference=ai_provider if ai_provider else None
                    )
                    
                    # Re-check length
                    result_word_count = len(result.split())
                    length_ratio = result_word_count / original_word_count
                    logger.info(f"After expansion: {result_word_count} words, ratio: {length_ratio:.2f}")
            
            return jsonify({'result': result})
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor for critique rewrite: {str(e)}")
            return jsonify({'error': f"Error processing critique rewrite: {str(e)}"}), 500
            
    except Exception as e:
        logger.error(f"Error in rewrite_from_output: {str(e)}")
        return jsonify({'error': f"Error processing critique rewrite: {str(e)}"}), 500

@app.route('/reset_api_keys', methods=['POST'])
def reset_api_keys():
    """Reset all API keys to available status"""
    try:
        # Force reload the API key manager to pick up current environment variables
        import importlib
        import api_key_manager
        importlib.reload(api_key_manager)
        from api_key_manager import api_key_manager as reloaded_manager
        
        # Reset all keys using the dedicated method
        reset_count = reloaded_manager.reset_all_keys()
        
        # Also reload the multi_provider_processor to use updated keys
        import multi_provider_processor
        importlib.reload(multi_provider_processor)
        
        # Log the action
        logger.info(f"API key reset requested. {reset_count} keys were reset to available status.")
        
        # Return success response
        return jsonify({
            'success': True, 
            'message': f'Successfully reactivated {reset_count} API keys.',
            'reset_count': reset_count
        })
    except Exception as e:
        # Log the error
        logger.error(f"Error resetting API keys: {str(e)}")
        
        # Return error response
        return jsonify({
            'success': False, 
            'message': f'Failed to reset API keys: {str(e)}'
        }), 500
        


@app.route('/share_text', methods=['POST'])
def share_text():
    """Share processed text via email using SendGrid"""
    try:
        data = request.get_json()
        recipient_email = data.get('email', '').strip()
        text_content = data.get('text', '').strip()
        subject = data.get('subject', 'Shared Processed Text')
        
        if not recipient_email:
            return jsonify({'error': 'Email address is required'}), 400
            
        if not text_content:
            return jsonify({'error': 'No text content to share'}), 400
            
        # Get SendGrid API key from environment
        sendgrid_api_key = os.environ.get('SENDGRID_API_KEY')
        verified_sender = 'jm@analyticphilosophy.ai'  # SendGrid verified sender
        
        if not sendgrid_api_key:
            return jsonify({'error': 'Email service not configured'}), 500
            
        # Import SendGrid modules
        from sendgrid import SendGridAPIClient
        from sendgrid.helpers.mail import Mail
        
        # Create email content
        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; margin: 20px;">
            <h2>Shared Text Content</h2>
            <div style="background-color: #f8f9fa; padding: 20px; border-radius: 5px; white-space: pre-wrap; font-family: monospace;">{text_content}</div>
            <p style="margin-top: 20px; color: #666; font-size: 12px;">
                This content was shared from an AI text processing application.
            </p>
        </body>
        </html>
        """
        
        # Create the email message
        message = Mail(
            from_email=verified_sender,
            to_emails=recipient_email,
            subject=subject,
            html_content=html_content
        )
        
        # Send the email
        sg = SendGridAPIClient(api_key=sendgrid_api_key)
        response = sg.send(message)
        
        if response.status_code in [200, 202]:
            return jsonify({'success': True, 'message': 'Text shared successfully!'})
        else:
            return jsonify({'error': f'Failed to send email. Status code: {response.status_code}'}), 500
            
    except Exception as e:
        app.logger.error(f"Error sharing text: {str(e)}")
        return jsonify({'error': f'Failed to share text: {str(e)}'}), 500

@app.route('/chat_with_ai', methods=['POST'])
def chat_with_ai():
    """Chat with AI that can see context from input/output boxes - FULL CONVERSATIONAL DIALOGUE"""
    try:
        data = request.get_json()
        user_message = data.get('message', '').strip()
        input_text = data.get('input_text', '').strip()
        output_text = data.get('output_text', '').strip()
        conversation_history = data.get('conversation_history', [])
        
        if not user_message:
            return jsonify({'error': 'Message is required'}), 400
        
        # Use OpenAI for full conversational capability
        import openai
        
        openai_key = os.environ.get('OPENAI_API_KEY')
        if not openai_key:
            return jsonify({'error': 'OpenAI API key not configured'}), 500
            
        openai_client = openai.OpenAI(api_key=openai_key)
        
        # Build system message with context awareness
        system_message = """You are a helpful AI assistant integrated into a text processing application called EZ Reader. You can help users with ANY questions, generate content, discuss topics, or assist with their work.

CAPABILITIES:
- Answer questions on any topic (like ChatGPT, Claude, or DeepSeek)
- Generate creative content, essays, stories, code, etc.
- Analyze and improve user's text from the input/output boxes
- Have natural, unlimited conversations on any subject
- Provide detailed explanations and reasoning

IMPORTANT: Users can send any of your responses to their input box for further processing. When generating content they might want to process, mention this capability."""

        # Add context if there's content in the boxes
        if input_text or output_text:
            context_parts = []
            if input_text:
                context_parts.append(f"INPUT BOX: {input_text[:500]}{'...' if len(input_text) > 500 else ''}")
            if output_text:
                context_parts.append(f"OUTPUT BOX: {output_text[:500]}{'...' if len(output_text) > 500 else ''}")
            system_message += f"\n\nCURRENT CONTEXT:\n" + "\n".join(context_parts)
        
        # Build messages array with conversation history
        messages = [{"role": "system", "content": system_message}]
        
        # Add conversation history
        for msg in conversation_history:
            messages.append({"role": msg['role'], "content": msg['content']})
        
        # Add current user message
        messages.append({"role": "user", "content": user_message})
        
        # Get AI response
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=2000,
            temperature=0.7
        )
        
        ai_response = response.choices[0].message.content
        
        return jsonify({
            'success': True,
            'response': ai_response
        })
            
    except Exception as e:
        app.logger.error(f"Error in chat: {str(e)}")
        return jsonify({'error': f'Chat error: {str(e)}'}), 500

@app.route('/homework_direct', methods=['POST'])
def homework_direct():
    """
    Process text directly as homework with streaming, chunked into 500-word outputs with 10-second pauses.
    """
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # CRITICAL: Remove all dollar signs from input
        text = text.replace('$', '')
        
        # Homework-specific prompt that handles ALL academic subjects
        homework_prompt = f"""CRITICAL: Do not use the $ symbol in your response. For math formatting, use \\( and \\) for inline math, and \\[ and \\] for display math.

You are an expert academic tutor specializing in all subjects including mathematics, philosophy, literature, history, science, essays, and any academic discipline. Complete the following assignment with thorough, detailed work:

{text}

Instructions:
1. Identify the subject area and type of assignment (math problem, essay question, philosophy prompt, etc.)
2. Provide complete, detailed responses appropriate to the academic level
3. For mathematics: Show ALL steps with proper LaTeX notation (use \\frac{{a}}{{b}}, x^{{2}}, \\sqrt{{x}}, etc.)
4. For essays/philosophy: Provide structured arguments with clear reasoning and examples
5. For literature: Include textual analysis and proper citations when relevant
6. For science: Explain concepts thoroughly with examples and applications
7. For any subject: Use proper academic formatting and terminology
8. Show your complete work and reasoning process
9. Use LaTeX notation for mathematical expressions: \\(...\\) for inline math, \\[...\\] for display math
10. IMPORTANT: Format your response with proper paragraphs separated by double line breaks (\\n\\n) for readability

Complete this assignment thoroughly and professionally:"""

        # Generate with automatic key + provider failover.
        import json
        from flask import stream_with_context
        from ai_failover import generate_with_failover

        @stream_with_context
        def generate():
            try:
                # Try every provider/key until one succeeds (OpenAI first for homework).
                result = generate_with_failover(
                    homework_prompt,
                    max_tokens=16000,
                    temperature=0.3,
                    preferred_order=['openai', 'anthropic', 'deepseek', 'azure', 'venice'],
                )
                full_response = result['text'].replace('$', '')

                # Strip Markdown so it renders cleanly in the plain-text output box
                cleaned_response = clean_markdown(full_response)

                # APPLY FORCED PARAGRAPH FORMATTING - GUARANTEED BREAKS EVERY 4 SENTENCES
                formatted_response = force_paragraph_formatting(cleaned_response)
                
                # Stream in chunks of ~50 characters. Each chunk is JSON-encoded so
                # newlines in the payload never break the SSE "\n\n" event framing.
                chunk_size = 50
                total_words = 0
                
                for i in range(0, len(formatted_response), chunk_size):
                    chunk = formatted_response[i:i+chunk_size]
                    yield f"data: {json.dumps({'text': chunk})}\n\n"
                    
                    # Count words for pause timing
                    total_words += len(chunk.split())
                    
                    # Pause every 500 words (control event - not shown to user)
                    if total_words >= 500:
                        yield f"data: {json.dumps({'pause': True})}\n\n"
                        time.sleep(10)
                        total_words = 0
                    
            except Exception as e:
                logger.error(f"Streaming error in homework: {str(e)}")
                yield f"data: {json.dumps({'text': chr(10) + chr(10) + 'Error: ' + str(e)})}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Homework direct processing error: {str(e)}")
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@app.route('/reconstruction/start_stream', methods=['POST'])
def reconstruction_start_stream():
    """Cross-Chunk Coherence (CC) long-document reconstruction.

    Three-pass architecture with Neon Postgres for intermediate state:
      PASS 1: Global skeleton extraction (thesis, outline, key terms, commitments)
      PASS 2: Constrained per-chunk processing with length enforcement
      PASS 3: Global consistency stitch + final assembly

    Streams Server-Sent Events with progress and final output text.
    """
    try:
        data = request.get_json(force=True) or {}
        original_text = (data.get('text') or '').strip()
        custom_instructions = (data.get('custom_instructions') or '').strip()
        document_title = data.get('document_title') or None

        if not original_text:
            return jsonify({'error': 'No text provided'}), 400
        if len(original_text.split()) < 100:
            return jsonify({'error': 'Long-document reconstruction requires at least 100 words. Use regular Rewrite for shorter texts.'}), 400

        user_id = 'anonymous'

        from flask import stream_with_context
        gen = run_reconstruction_stream(original_text, custom_instructions, user_id, document_title)

        return Response(stream_with_context(gen), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive',
        })
    except Exception as e:
        logger.exception("reconstruction_start_stream failed")
        return jsonify({'error': f'Reconstruction failed: {str(e)}'}), 500


@app.route('/reconstruction/status/<job_id>', methods=['GET'])
def reconstruction_status(job_id):
    """Inspect a reconstruction job's current state in Neon."""
    try:
        from reconstruction_engine import _exec
        job = _exec(
            "SELECT id, status, current_chunk, num_chunks, total_input_words, "
            "target_min_words, target_max_words, length_mode, final_word_count, "
            "error_message, created_at, updated_at FROM reconstruction_jobs WHERE id=%s",
            (job_id,), fetch="one")
        if not job:
            return jsonify({'error': 'Job not found'}), 404
        return jsonify({k: (str(v) if hasattr(v, 'isoformat') or k == 'id' else v) for k, v in job.items()})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/reconstruction/result/<job_id>', methods=['GET'])
def reconstruction_result(job_id):
    """Fetch final assembled output for a completed reconstruction job."""
    try:
        from reconstruction_engine import _exec
        job = _exec("SELECT status, final_output, final_word_count FROM reconstruction_jobs WHERE id=%s",
                    (job_id,), fetch="one")
        if not job:
            return jsonify({'error': 'Job not found'}), 404
        return jsonify({
            'status': job['status'],
            'final_output': job['final_output'] or '',
            'final_word_count': job['final_word_count'] or 0,
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/style_rewrite_passthrough', methods=['POST'])
def style_rewrite_passthrough_route():
    """
    Process a style rewrite request using the pure pass-through mechanism.
    This endpoint accepts a style sample and target text, and passes them directly
    to the style_rewrite_passthrough module without any modifications.
    
    Request JSON parameters:
    - style_sample: Text representing the user's writing style
    - target_text: Text to be rewritten in the user's style
    
    Returns:
    - JSON response with the rewritten text
    """
    try:
        data = request.get_json()
        style_sample = data.get('style_sample', '')
        target_text = data.get('target_text', '')
        
        if not style_sample or not target_text:
            return jsonify({'error': 'Both style sample and target text are required'}), 400
            
        # Log length of inputs for debugging
        logger.debug(f"Style sample length: {len(style_sample)} chars")
        logger.debug(f"Target text length: {len(target_text)} chars")
        
        # Call the passthrough style rewrite function
        result = process_style_rewrite(style_sample, target_text)
        
        return jsonify({
            'success': True,
            'result': result
        })
        
    except Exception as e:
        logger.error(f"Error in style rewrite passthrough: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/comprehensive_search', methods=['POST'])
def comprehensive_search():
    """
    Perform comprehensive search using Google CSE and multi-AI research.
    Automatically extracts search terms from text content if no query provided.
    """
    try:
        data = request.json
        query = data.get('query', '').strip()
        text_content = data.get('text_content', '').strip()
        
        # Import the comprehensive search module
        from comprehensive_search import perform_comprehensive_search
        
        # Perform the search
        results = perform_comprehensive_search(query=query if query else None, text_content=text_content)
        
        return jsonify({
            'success': True,
            'results': results
        })
        
    except Exception as e:
        logger.error(f"Comprehensive search error: {e}")
        return jsonify({
            'success': False,
            'error': f'Search failed: {str(e)}'
        }), 500

@app.route('/create_podcast', methods=['POST'])
def create_podcast():
    """
    Create a professional podcast from text content with OpenAI voices
    Supports single host, dual host, and interview modes
    """
    try:
        from podcast_generator import PodcastGenerator
        
        data = request.json
        content = data.get('content', '').strip()
        mode = data.get('mode', 'dual_host')  # single_host, dual_host, interview
        host1_voice = data.get('host1_voice', 'alloy')
        host2_voice = data.get('host2_voice', 'echo')
        host1_name = data.get('host1_name', 'Alex')
        host2_name = data.get('host2_name', 'Sam')
        style = data.get('style', 'conversational')
        
        if not content:
            return jsonify({'error': 'No content provided for podcast creation'}), 400
        
        # Preprocess dollar signs from content
        content = preprocess_dollar_signs(content)
        
        # Initialize podcast generator
        generator = PodcastGenerator()
        
        # Generate unique filename
        timestamp = int(time.time())
        unique_id = str(uuid.uuid4())[:8]
        output_filename = f"podcast_{timestamp}_{unique_id}.mp3"
        output_path = os.path.join('uploads', output_filename)
        
        # Ensure uploads directory exists
        os.makedirs('uploads', exist_ok=True)
        
        # Create podcast
        success, result = generator.create_podcast(
            content=content,
            output_path=output_path,
            mode=mode,
            host1_voice=host1_voice,
            host2_voice=host2_voice,
            host1_name=host1_name,
            host2_name=host2_name,
            style=style
        )
        
        if success:
            return jsonify({
                'success': True,
                'audio_url': f'/uploads/{output_filename}',
                'filename': output_filename,
                'metadata': result if isinstance(result, dict) else {
                    'path': output_path,
                    'mode': mode,
                    'hosts': [host1_name, host2_name] if mode != 'single_host' else [host1_name]
                }
            })
        else:
            return jsonify({
                'success': False,
                'error': result
            }), 500
            
    except Exception as e:
        logger.error(f"Error creating podcast: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Podcast creation failed: {str(e)}'
        }), 500

@app.route('/get_podcast_voices', methods=['GET'])
def get_podcast_voices():
    """
    Get available OpenAI TTS voices for podcast creation
    """
    try:
        from podcast_generator import PodcastGenerator
        
        generator = PodcastGenerator()
        voices = generator.get_available_voices()
        
        return jsonify({
            'success': True,
            'voices': voices
        })
        
    except Exception as e:
        logger.error(f"Error getting podcast voices: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/download/<filename>')
def download_file(filename):
    """Download audio files from static/audio directory"""
    return send_from_directory('static/audio', filename, as_attachment=True)


@app.route('/quality_assessment_stream', methods=['POST'])
def quality_assessment_stream():
    """
    Quality assessment using Anthropic with streaming response and chunking for large texts.
    """
    try:
        from anthropic import Anthropic
        import os
        import re
        
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Get Anthropic API key
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
        if not anthropic_key:
            return jsonify({'error': 'Anthropic API key not configured'}), 500
        
        # Initialize Anthropic client
        client = Anthropic(api_key=anthropic_key)
        
        # Chunk text if it's too large (split at ~2000 words per chunk to avoid token limits)
        words = text.split()
        word_count = len(words)
        chunks = []
        
        if word_count > 2000:
            # Split into chunks at paragraph boundaries
            paragraphs = re.split(r'\n\s*\n', text)
            current_chunk = []
            current_word_count = 0
            
            for para in paragraphs:
                para_words = len(para.split())
                if current_word_count + para_words > 2000 and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_word_count = para_words
                else:
                    current_chunk.append(para)
                    current_word_count += para_words
            
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
        else:
            chunks = [text]
        
        # Build the exact prompt as specified
        def create_prompt(chunk_text, chunk_num=None, total_chunks=None):
            chunk_info = f" (CHUNK {chunk_num} OF {total_chunks})" if chunk_num else ""
            return f"""Here is the text to analyze{chunk_info}:

{chunk_text}

---

Now answer these questions about the text above.

FORMATTING REQUIREMENTS (MANDATORY):
- Use clear headings for each question (use ## for markdown headers)
- Add blank lines between sections for readability
- Use paragraphs - DO NOT write one giant block of text
- Include a score in N/100 format for EVERY question (e.g., "Score: 73/100")
- End with an OVERALL SCORE that averages/synthesizes all scores

DEPTH REQUIREMENTS (CRITICAL):
- This must be a THOROUGH, COMPREHENSIVE, EXTENSIVE analysis - NOT a quick summary
- Write AT LEAST 4-6 substantial paragraphs for EACH question
- Quote extensively and analyze each quote in detail
- Reference multiple specific passages for each question
- Develop your analysis deeply - explore nuances, implications, and connections
- Take your time to provide genuine intellectual analysis, not surface-level observations
- This should be a LONG, DETAILED response that demonstrates deep engagement with the text

ANTI-TRUNCATION REQUIREMENTS (ABSOLUTELY MANDATORY):
- NEVER truncate your analysis
- NEVER add disclaimers like "I can provide more if desired" or "condensed due to space limitations"
- NEVER claim space constraints prevent full analysis
- NEVER ask the user if they want you to continue or proceed with the full analysis
- NEVER stop mid-analysis to ask for permission to continue
- You MUST provide COMPLETE analysis for ALL sections WITHOUT ASKING
- There are NO space limitations - provide the FULL analysis AUTOMATICALLY
- Do NOT abbreviate or summarize - give the COMPLETE, THOROUGH analysis for every single question
- Just DO the complete analysis - don't ask, don't wait for permission, don't pause

CRITICAL: You MUST quote EXTENSIVELY and SPECIFICALLY from the text. Every answer must include DIRECT QUOTES and SPECIFIC REFERENCES to the text. Do NOT give generic responses that could apply to any text.

## IS IT INSIGHTFUL?
(Provide score: X/100)

## DOES IT DEVELOP POINTS? (OR, IF IT IS A SHORT EXCERPT, IS THERE EVIDENCE THAT IT WOULD DEVELOP POINTS IF EXTENDED)?
(Provide score: X/100)

## IS THE ORGANIZATION MERELY SEQUENTIAL (JUST ONE POINT AFTER ANOTHER, LITTLE OR NO LOGICAL SCAFFOLDING)? OR ARE THE IDEAS ARRANGED, NOT JUST SEQUENTIALLY BUT HIERARCHICALLY?
(Provide score: X/100)

## IF THE POINTS IT MAKES ARE NOT INSIGHTFUL, DOES IT OPERATE SKILLFULLY WITH CANONS OF LOGIC/REASONING?
(Provide score: X/100)

## ARE THE POINTS CLICHES? OR ARE THEY "FRESH"?
(Provide score: X/100)

## DOES IT USE TECHNICAL JARGON TO OBFUSCATE OR TO RENDER MORE PRECISE?
(Provide score: X/100)

## IS IT ORGANIC? DO POINTS DEVELOP IN AN ORGANIC, NATURAL WAY? DO THEY 'UNFOLD'? OR ARE THEY FORCED AND ARTIFICIAL?
(Provide score: X/100)

## DOES IT OPEN UP NEW DOMAINS? OR, ON THE CONTRARY, DOES IT SHUT OFF INQUIRY (BY CONDITIONALIZING FURTHER DISCUSSION OF THE MATTERS ON ACCEPTANCE OF ITS INTERNAL AND POSSIBLY VERY FAULTY LOGIC)?
(Provide score: X/100)

## IS IT ACTUALLY INTELLIGENT OR JUST THE WORK OF SOMEBODY WHO, JUDGING BY THE SUBJECT-MATTER, IS PRESUMED TO BE INTELLIGENT (BUT MAY NOT BE)?
(Provide score: X/100)

## IS IT REAL OR IS IT PHONY?
(Provide score: X/100)

## DO THE SENTENCES EXHIBIT COMPLEX AND COHERENT INTERNAL LOGIC?
(Provide score: X/100)

## IS THE PASSAGE GOVERNED BY A STRONG CONCEPT? OR IS THE ONLY ORGANIZATION DRIVEN PURELY BY EXPOSITORY (AS OPPOSED TO EPISTEMIC) NORMS?
(Provide score: X/100)

## IS THERE SYSTEM-LEVEL CONTROL OVER IDEAS? IN OTHER WORDS, DOES THE AUTHOR SEEM TO RECALL WHAT HE SAID EARLIER AND TO BE IN A POSITION TO INTEGRATE IT INTO POINTS HE HAS MADE SINCE THEN?
(Provide score: X/100)

## ARE THE POINTS 'REAL'? ARE THEY FRESH? OR IS SOME INSTITUTION OR SOME ACCEPTED VEIN OF PROPAGANDA OR ORTHODOXY JUST USING THE AUTHOR AS A MOUTH PIECE?
(Provide score: X/100)

## IS THE WRITING EVASIVE OR DIRECT?
(Provide score: X/100)

## ARE THE STATEMENTS AMBIGUOUS?
(Provide score: X/100)

## DOES THE PROGRESSION OF THE TEXT DEVELOP ACCORDING TO WHO SAID WHAT OR ACCORDING TO WHAT ENTAILS OR CONFIRMS WHAT?
(Provide score: X/100)

## DOES THE AUTHOR USE OTHER AUTHORS TO DEVELOP HIS IDEAS OR TO CLOAK HIS OWN LACK OF IDEAS?
(Provide score: X/100)

## OVERALL SCORE
(Provide overall score: X/100, representing composite assessment)

CRITICAL INSTRUCTIONS:
- You MUST engage deeply with the actual text. Quote specific passages extensively.
- You MUST reference specific arguments, distinctions, and examples from the text.
- You MUST demonstrate that you have read and understood the text by pointing to concrete textual evidence.
- DO NOT give superficial, generic answers that could apply to any text.
- Each answer must be grounded in the actual content and arguments of this specific text.
- MUST include a score for EVERY question and an overall score at the end.

IMPORTANT QUALIFICATIONS:
- A score of N/100 (e.g. 73/100) means that (100-N)/100 (e.g. 27/100) of writers outperform the author with respect to the parameter defined by the question.
- You are NOT grading; you are answering these questions.
- You do NOT use a risk-averse standard; you do NOT attempt to be diplomatic; you do NOT attempt to comply with risk-averse, medium-range IQ, academic norms.
- You do NOT make assumptions about the level of the paper; it could be a work of the highest excellence and genius, or it could be the work of a moron.
- If a work is a work of genius, you SAY that, and you SAY WHY; you do NOT shy away from giving what might conventionally be regarded as excessively "superlative" scores; you give it the score it deserves, not the score that a midwit committee would say it deserves.
- Think VERY VERY VERY hard about your answers; do NOT default to cookbook, midwit evaluation protocols.
- DO NOT give credit merely for use of jargon or for referencing authorities. Focus on SUBSTANCE. Only give points for scholarly references/jargon if they unambiguously increase substance."""
        
        # Create streaming response with chunking support
        from flask import stream_with_context
        import time
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                
                for i, chunk in enumerate(chunks, 1):
                    if total_chunks > 1:
                        chunk_header = f"\n\n{'='*50}\nANALYSIS OF CHUNK {i} OF {total_chunks}\n{'='*50}\n\n"
                        yield f"data: {chunk_header}\n\n"
                    
                    prompt = create_prompt(chunk, i if total_chunks > 1 else None, total_chunks if total_chunks > 1 else None)
                    
                    with client.messages.stream(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=6000,
                        system="You are an expert text analyst. When given text to analyze, you MUST analyze it directly without asking for confirmation or claiming the text wasn't provided. Analyze the text immediately and thoroughly. NEVER truncate your analysis or add disclaimers about space limitations. NEVER ask the user if they want you to continue - just provide the COMPLETE, FULL analysis for ALL sections WITHOUT ASKING. Do the entire analysis automatically.",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3
                    ) as stream:
                        for text_chunk in stream.text_stream:
                            # Send each chunk as SSE format
                            yield f"data: {text_chunk}\n\n"
                    
                    # Add delay between chunks to avoid rate limits
                    if i < total_chunks:
                        yield f"data: \n\n\n\n"
                        yield f"data: [Pausing 3 seconds to avoid rate limits...]\n\n"
                        time.sleep(3)  # 3 second delay between chunks
                        yield f"data: \n\n"
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in quality assessment: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/quality_writing_assessment_stream', methods=['POST'])
def quality_writing_assessment_stream():
    """
    Quality writing assessment using Anthropic with streaming response and chunking for large texts.
    """
    try:
        from anthropic import Anthropic
        import os
        import re
        
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Get Anthropic API key
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
        if not anthropic_key:
            return jsonify({'error': 'Anthropic API key not configured'}), 500
        
        # Initialize Anthropic client
        client = Anthropic(api_key=anthropic_key)
        
        # Chunk text if it's too large (split at ~2000 words per chunk to avoid token limits)
        words = text.split()
        word_count = len(words)
        chunks = []
        
        if word_count > 2000:
            # Split into chunks at paragraph boundaries
            paragraphs = re.split(r'\n\s*\n', text)
            current_chunk = []
            current_word_count = 0
            
            for para in paragraphs:
                para_words = len(para.split())
                if current_word_count + para_words > 2000 and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_word_count = para_words
                else:
                    current_chunk.append(para)
                    current_word_count += para_words
            
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
        else:
            chunks = [text]
        
        # Build the quality assessment prompt
        def create_quality_prompt(chunk_text, chunk_num=None, total_chunks=None):
            chunk_info = f" (CHUNK {chunk_num} OF {total_chunks})" if chunk_num else ""
            return f"""Here is the text to analyze{chunk_info}:

{chunk_text}

---

Now answer these questions about the text above.

FORMATTING REQUIREMENTS (MANDATORY):
- Use clear headings for each section and question (use ## for sections, ### for questions)
- Add blank lines between sections for readability
- Use paragraphs - DO NOT write one giant block of text
- Provide thoughtful analysis for each question with specific textual evidence
- End with an OVERALL QUALITY SCORE (1-100)

DEPTH REQUIREMENTS:
- Provide thorough, comprehensive analysis for each question
- Quote extensively from the text to support your assessments
- Reference specific passages and examples
- Develop your analysis deeply - explore nuances and implications
- This should be a LONG, DETAILED response demonstrating deep engagement

ANTI-TRUNCATION REQUIREMENTS (ABSOLUTELY MANDATORY):
- NEVER truncate your analysis
- NEVER add disclaimers like "I can provide more if desired" or "condensed due to space limitations"
- NEVER claim space constraints prevent full analysis
- NEVER ask the user if they want you to continue or proceed with the full analysis
- NEVER stop mid-analysis to ask for permission to continue
- You MUST provide COMPLETE analysis for ALL sections and ALL questions WITHOUT ASKING
- There are NO space limitations - provide the FULL analysis AUTOMATICALLY
- Do NOT abbreviate or summarize - give the COMPLETE, THOROUGH analysis for every single question
- Just DO the complete analysis - don't ask, don't wait for permission, don't pause

## Structural and Stylistic Quality

### Does the text maintain a coherent structure from start to finish?

### Are transitions between sections smooth and logically motivated?

### Does each paragraph contribute clearly to the central argument or theme?

### Is the pacing appropriate — neither rushed nor meandering?

### Does the writing exhibit rhythm and stylistic control at the sentence level?

## Clarity and Readability

### Are key concepts and terms clearly defined and consistently used?

### Does the text communicate complex ideas without unnecessary jargon?

### Is there an optimal balance between concision and elaboration?

### Does the prose flow naturally, or does it feel stilted or mechanical?

### Is the argument easy to follow without external guidance?

## Substantive and Evidential Quality

### Are claims supported by adequate reasoning, examples, or data?

### Does the evidence genuinely connect to the claims it's meant to support?

### Are counterarguments acknowledged or integrated where appropriate?

### Is there an appropriate balance between generalization and specificity?

### Does the text display awareness of its own limitations or assumptions?

## Voice and Engagement

### Does the author's voice convey confidence and authority without bluster?

### Does the tone match the subject matter and intended audience?

### Is there stylistic distinctiveness, or does it feel generic or derivative?

### Does the prose sustain reader engagement across its full length?

### Does it exhibit emotional or rhetorical intelligence — i.e., timing, emphasis, and restraint?

## Empirical / Historical Quality

### Are factual claims verifiable and consistent with established data or sources?

### Does the author distinguish clearly between evidence, interpretation, and speculation?

### Are historical or experimental contexts accurately represented, not retrofitted to the thesis?

### Does the argument derive naturally from the data, rather than cherry-picking it?

### Are cause–effect relations supported rather than merely asserted?

### Is statistical or quantitative reasoning (if used) presented correctly and transparently?

### Does the essay acknowledge alternative explanations or competing evidence?

### Are timelines, causal chains, and data sequences coherent and non-contradictory?

### Does the writer show an understanding of methodological limits or uncertainty?

### Does the synthesis of evidence reveal genuine insight, not just accumulation of facts?

## Logical / Philosophical / Mathematical Quality

### Are all key terms, variables, or premises defined precisely and used consistently?

### Are inferences valid — that is, do conclusions actually follow from premises?

### Are distinctions (e.g., necessary vs. sufficient, analytic vs. synthetic) properly maintained?

### Does the essay avoid equivocation and circular reasoning?

### Are counterexamples anticipated and dealt with rigorously?

### Is the progression of ideas deductively or dialectically coherent (no leaps or gaps)?

### Are formal or quasi-formal structures (proofs, schemata, analogies) accurately executed?

### Does the argument integrate examples or thought-experiments that genuinely test its claims?

### Are theoretical claims connected to real conceptual or empirical stakes, not pure wordplay?

### Is the reasoning self-monitoring — aware of its own assumptions, scope, and implications?

## OVERALL QUALITY SCORE
(Provide overall quality score: X/100, representing composite assessment across all dimensions)

CRITICAL INSTRUCTIONS:
- You MUST engage deeply with the actual text
- Quote specific passages extensively to support your assessments
- Reference specific arguments, distinctions, and examples from the text
- Each answer must be grounded in the actual content of this specific text
- DO NOT give superficial, generic answers"""
        
        # Create streaming response with chunking support
        from flask import stream_with_context
        import time
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                
                for i, chunk in enumerate(chunks, 1):
                    if total_chunks > 1:
                        chunk_header = f"\n\n{'='*50}\nQUALITY ANALYSIS OF CHUNK {i} OF {total_chunks}\n{'='*50}\n\n"
                        yield f"data: {chunk_header}\n\n"
                    
                    prompt = create_quality_prompt(chunk, i if total_chunks > 1 else None, total_chunks if total_chunks > 1 else None)
                    
                    with client.messages.stream(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=6000,
                        system="You are an expert writing quality analyst. When given text to analyze, you MUST analyze it directly and thoroughly using the provided framework. Provide detailed, specific assessments grounded in the actual text. NEVER truncate your analysis or add disclaimers about space limitations. NEVER ask the user if they want you to continue - just provide the COMPLETE, FULL analysis for ALL sections and ALL questions WITHOUT ASKING. Do the entire analysis automatically.",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3
                    ) as stream:
                        for text_chunk in stream.text_stream:
                            # Send each chunk as SSE format
                            yield f"data: {text_chunk}\n\n"
                    
                    # Add delay between chunks to avoid rate limits
                    if i < total_chunks:
                        yield f"data: \n\n\n\n"
                        yield f"data: [Pausing 3 seconds to avoid rate limits...]\n\n"
                        time.sleep(3)  # 3 second delay between chunks
                        yield f"data: \n\n"
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in quality writing assessment: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/fiction_assessment_stream', methods=['POST'])
def fiction_assessment_stream():
    """
    Fiction quality assessment using Anthropic with streaming response and chunking for large texts.
    """
    try:
        from anthropic import Anthropic
        import os
        import re
        
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Get Anthropic API key
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
        if not anthropic_key:
            return jsonify({'error': 'Anthropic API key not configured'}), 500
        
        # Initialize Anthropic client
        client = Anthropic(api_key=anthropic_key)
        
        # Chunk text if it's too large (split at ~2000 words per chunk to avoid token limits)
        words = text.split()
        word_count = len(words)
        chunks = []
        
        if word_count > 2000:
            # Split into chunks at paragraph boundaries
            paragraphs = re.split(r'\n\s*\n', text)
            current_chunk = []
            current_word_count = 0
            
            for para in paragraphs:
                para_words = len(para.split())
                if current_word_count + para_words > 2000 and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_word_count = para_words
                else:
                    current_chunk.append(para)
                    current_word_count += para_words
            
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
        else:
            chunks = [text]
        
        # Build the fiction assessment prompt
        def create_fiction_prompt(chunk_text, chunk_num=None, total_chunks=None):
            chunk_info = f" (CHUNK {chunk_num} OF {total_chunks})" if chunk_num else ""
            return f"""Here is the fiction text to analyze{chunk_info}:

{chunk_text}

---

Now answer these questions about the fiction text above.

FORMATTING REQUIREMENTS (MANDATORY):
- Use clear headings for each section and question (use ## for sections, ### for questions)
- Add blank lines between sections for readability
- Use paragraphs - DO NOT write one giant block of text
- Provide thoughtful analysis for each question with specific textual evidence
- End with an OVERALL FICTION QUALITY SCORE (1-100)

DEPTH REQUIREMENTS:
- Provide thorough, comprehensive analysis for each question
- Quote extensively from the text to support your assessments
- Reference specific passages and examples
- Develop your analysis deeply - explore nuances and implications
- This should be a LONG, DETAILED response demonstrating deep engagement

ANTI-TRUNCATION REQUIREMENTS (ABSOLUTELY MANDATORY):
- NEVER truncate your analysis
- NEVER add disclaimers like "I can provide more if desired" or "condensed due to space limitations"
- NEVER claim space constraints prevent full analysis
- NEVER ask the user if they want you to continue or proceed with the full analysis
- NEVER stop mid-analysis to ask for permission to continue
- You MUST provide COMPLETE analysis for ALL sections and ALL questions WITHOUT ASKING
- There are NO space limitations - provide the FULL analysis AUTOMATICALLY
- Do NOT abbreviate or summarize - give the COMPLETE, THOROUGH analysis for every single question
- Just DO the complete analysis - don't ask, don't wait for permission, don't pause

## A. Narrative Structure and Coherence

### Does the story have a clear structural arc (setup, conflict, resolution)?

### Are plot events causally connected rather than random or episodic?

### Does the pacing sustain tension and momentum?

### Are scene transitions smooth and meaningful?

### Is there narrative unity — do all scenes serve the central story or theme?

### Does the story begin and end at the right points?

### Is exposition handled naturally, without info-dumps?

### Does the narrative escalate stakes intelligently rather than through contrivance?

### Are time shifts (flashbacks, nonlinear structure) handled clearly?

### Is the story's point of view (first, third, omniscient, etc.) consistently maintained?

## B. Characterization and Psychology

### Are characters psychologically believable and internally consistent?

### Do they reveal themselves through action and dialogue rather than exposition?

### Are motivations clear without being simplistic?

### Do main characters change or deepen over the course of the story?

### Are secondary characters distinct and purposeful, not filler?

### Is there moral or psychological ambiguity where appropriate?

### Are character relationships developed with nuance and tension?

### Do inner conflicts reflect larger thematic tensions?

### Are characters' speech patterns and thoughts suited to their background and situation?

### Do the characters' choices drive the plot rather than convenience or coincidence?

## C. Style and Language

### Does the prose show rhythm, texture, and tonal control?

### Are images vivid without being overwrought or clichéd?

### Does the diction fit the narrative voice and setting?

### Is figurative language (metaphor, simile, symbol) purposeful, not decorative?

### Does sentence structure vary to reflect emotional shifts or pacing?

### Does dialogue sound natural and reveal character?

### Is narration precise — neither vague nor florid?

### Does the text have a distinctive voice or stylistic fingerprint?

### Is the tone internally coherent (comic, tragic, lyrical, ironic, etc.)?

### Does the language evoke sensory and emotional immediacy?

## D. Thematic and Conceptual Depth

### Does the story explore meaningful ideas or human truths?

### Are its themes integrated into the narrative rather than imposed from outside?

### Does it offer insight into psychology, society, or morality beyond the surface plot?

### Are symbols or motifs developed consistently and intelligently?

### Does the story resist trivial moralizing or cliché resolution?

### Is there interpretive richness — multiple plausible readings?

### Do the conflicts reflect deeper philosophical or existential questions?

### Is the theme embodied through character and action, not merely stated?

### Does the narrative invite reflection without forcing it?

### Does the story achieve closure without reducing complexity?

## E. Emotional and Aesthetic Impact

### Does the story evoke genuine emotional engagement?

### Are emotional beats earned through character development and situation?

### Does it balance intellect and feeling?

### Is there a rhythm of tension and release that sustains attention?

### Does it linger in the mind — memorable scenes, lines, or insights?

### Does the story feel alive — unpredictable yet inevitable?

### Does it leave the reader altered, challenged, or moved?

### Is its world immersive and internally consistent?

### Does it create aesthetic pleasure beyond plot satisfaction?

### Does it achieve resonance — a sense of meaning beyond the literal?

## OVERALL FICTION QUALITY SCORE
(Provide overall fiction quality score: X/100, representing composite assessment across all dimensions)

CRITICAL INSTRUCTIONS:
- You MUST engage deeply with the actual fiction text
- Quote specific passages extensively to support your assessments
- Reference specific characters, scenes, dialogue, and narrative moments from the text
- Each answer must be grounded in the actual content of this specific fiction work
- DO NOT give superficial, generic answers"""
        
        # Create streaming response with chunking support
        from flask import stream_with_context
        import time
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                
                for i, chunk in enumerate(chunks, 1):
                    if total_chunks > 1:
                        chunk_header = f"\n\n{'='*50}\nFICTION ANALYSIS OF CHUNK {i} OF {total_chunks}\n{'='*50}\n\n"
                        yield f"data: {chunk_header}\n\n"
                    
                    prompt = create_fiction_prompt(chunk, i if total_chunks > 1 else None, total_chunks if total_chunks > 1 else None)
                    
                    with client.messages.stream(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=6000,
                        system="You are an expert fiction analyst and literary critic. When given fiction text to analyze, you MUST analyze it directly and thoroughly using the provided framework. Provide detailed, specific assessments grounded in the actual text. NEVER truncate your analysis or add disclaimers about space limitations. NEVER ask the user if they want you to continue - just provide the COMPLETE, FULL analysis for ALL sections and ALL questions WITHOUT ASKING. Do the entire analysis automatically.",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3
                    ) as stream:
                        for text_chunk in stream.text_stream:
                            # Send each chunk as SSE format
                            yield f"data: {text_chunk}\n\n"
                    
                    # Add delay between chunks to avoid rate limits
                    if i < total_chunks:
                        yield f"data: \n\n\n\n"
                        yield f"data: [Pausing 3 seconds to avoid rate limits...]\n\n"
                        time.sleep(3)  # 3 second delay between chunks
                        yield f"data: \n\n"
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in fiction assessment: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/intelligence_maximization_stream', methods=['POST'])
def intelligence_maximization_stream():
    """
    Intelligence Maximization - Rewrites text to maximize intelligence scores with streaming response and chunking for large texts.
    """
    try:
        from anthropic import Anthropic
        import os
        import re
        
        data = request.get_json()
        text = data.get('text', '')
        assessment_results = data.get('assessment_results', '')  # Optional: previous assessment results
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Get Anthropic API key
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
        if not anthropic_key:
            return jsonify({'error': 'Anthropic API key not configured'}), 500
        
        # Initialize Anthropic client
        client = Anthropic(api_key=anthropic_key)
        
        # Chunk text if it's too large (split at ~2000 words per chunk to avoid token limits)
        words = text.split()
        word_count = len(words)
        chunks = []
        
        if word_count > 2000:
            # Split into chunks at paragraph boundaries
            paragraphs = re.split(r'\n\s*\n', text)
            current_chunk = []
            current_word_count = 0
            
            for para in paragraphs:
                para_words = len(para.split())
                if current_word_count + para_words > 2000 and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_word_count = para_words
                else:
                    current_chunk.append(para)
                    current_word_count += para_words
            
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
        else:
            chunks = [text]
        
        # Build the intelligence maximization prompt
        def create_maximization_prompt(chunk_text, chunk_num=None, total_chunks=None, has_assessment=False):
            chunk_info = f" (CHUNK {chunk_num} OF {total_chunks})" if chunk_num else ""
            
            assessment_context = ""
            if has_assessment and assessment_results:
                assessment_context = f"\n\nPREVIOUS INTELLIGENCE ASSESSMENT RESULTS:\n{assessment_results}\n\nUse these assessment results to guide your rewrite and address any weaknesses identified.\n"
            
            return f"""Here is the text to rewrite{chunk_info}:

{chunk_text}

---

You are tasked with rewriting the text above to MAXIMIZE its intelligence across all dimensions. The rewritten text should score maximally high on these criteria:

INTELLIGENCE MAXIMIZATION CRITERIA:

1. **INSIGHT & DEPTH**: Make it profoundly insightful. Replace superficial observations with deep, original insights.

2. **DEVELOPMENT OF POINTS**: Develop all points fully and systematically. If it's a short excerpt, ensure it shows clear potential for extended development.

3. **HIERARCHICAL ORGANIZATION**: Arrange ideas not just sequentially but hierarchically. Create a logical scaffolding where ideas build upon each other in sophisticated ways.

4. **LOGICAL/REASONING SKILL**: Demonstrate masterful use of logic and reasoning. Every inference should be valid and rigorous.

5. **FRESHNESS & ORIGINALITY**: Eliminate all clichés. Every point should feel fresh, original, and non-obvious.

6. **PRECISION NOT OBFUSCATION**: Use technical language only to achieve precision, never to obscure. Every term should clarify.

7. **ORGANIC DEVELOPMENT**: Points should unfold naturally and organically, not feel forced or artificial. Ideas should grow from each other.

8. **OPENS NEW DOMAINS**: The text should open up inquiry, not shut it off. Avoid conditional logic that constrains further discussion.

9. **GENUINE INTELLIGENCE**: Display actual intelligence in the content itself, not reliance on presumed expertise based on subject matter.

10. **AUTHENTICITY**: Make it real, not phony. Genuine insight over performance of intelligence.

11. **COMPLEX SENTENCE LOGIC**: Each sentence should exhibit complex, coherent internal logic.

12. **STRONG GOVERNING CONCEPT**: Organize the text around a strong central concept, not just expository norms.

13. **SYSTEM-LEVEL CONTROL**: Demonstrate control over the entire idea system. Integrate earlier points with later developments.

14. **FRESH REAL POINTS**: Every point should be real and fresh, not institutional propaganda or orthodoxy speaking through the text.

15. **DIRECTNESS NOT EVASION**: Be direct, not evasive. No ambiguity or hedging.

16. **LOGICAL PROGRESSION**: Progress should be driven by what entails or confirms what, not by "who said what."

17. **IDEAS NOT NAME-DROPPING**: Use other authors/sources to develop ideas, not to mask lack of ideas.
{assessment_context}
INSTRUCTIONS FOR REWRITING:
- Completely rewrite the text to maximize intelligence on ALL criteria above
- Preserve the core ideas but elevate them dramatically
- Make every sentence count - each should demonstrate sophisticated thought
- Create hierarchical logical structure with clear conceptual scaffolding
- Be bold and original - transform superficial points into profound insights
- Maintain or exceed the original length while adding depth, not padding
- Write in a voice that is confident, direct, and authentically intelligent

NOW PROVIDE THE MAXIMALLY INTELLIGENT REWRITE:"""
        
        # Create streaming response with chunking support
        from flask import stream_with_context
        import time
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                has_assessment = bool(assessment_results.strip())
                
                for i, chunk in enumerate(chunks, 1):
                    if total_chunks > 1:
                        chunk_header = f"\n\n{'='*50}\nINTELLIGENCE-MAXIMIZED REWRITE - CHUNK {i} OF {total_chunks}\n{'='*50}\n\n"
                        yield f"data: {chunk_header}\n\n"
                    
                    prompt = create_maximization_prompt(chunk, i if total_chunks > 1 else None, total_chunks if total_chunks > 1 else None, has_assessment)
                    
                    with client.messages.stream(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=6000,
                        system="You are an expert intellectual writer capable of transforming ordinary text into deeply intelligent, insightful prose. You rewrite to maximize insight, logical sophistication, originality, and organic development of ideas. You never produce superficial or clichéd content.",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7
                    ) as stream:
                        for text_chunk in stream.text_stream:
                            # Send each chunk as SSE format
                            yield f"data: {text_chunk}\n\n"
                    
                    # Add delay between chunks to avoid rate limits
                    if i < total_chunks:
                        yield f"data: \n\n\n\n"
                        yield f"data: [Pausing 3 seconds to avoid rate limits...]\n\n"
                        time.sleep(3)  # 3 second delay between chunks
                        yield f"data: \n\n"
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in intelligence maximization: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/quality_maximization_stream', methods=['POST'])
def quality_maximization_stream():
    """
    Quality Maximization - Rewrites text to maximize quality scores with streaming response and chunking for large texts.
    """
    try:
        from anthropic import Anthropic
        import os
        import re
        
        data = request.get_json()
        text = data.get('text', '')
        assessment_results = data.get('assessment_results', '')  # Optional: previous assessment results
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Get Anthropic API key
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
        if not anthropic_key:
            return jsonify({'error': 'Anthropic API key not configured'}), 500
        
        # Initialize Anthropic client
        client = Anthropic(api_key=anthropic_key)
        
        # Chunk text if it's too large (split at ~2000 words per chunk to avoid token limits)
        words = text.split()
        word_count = len(words)
        chunks = []
        
        if word_count > 2000:
            # Split into chunks at paragraph boundaries
            paragraphs = re.split(r'\n\s*\n', text)
            current_chunk = []
            current_word_count = 0
            
            for para in paragraphs:
                para_words = len(para.split())
                if current_word_count + para_words > 2000 and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_word_count = para_words
                else:
                    current_chunk.append(para)
                    current_word_count += para_words
            
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
        else:
            chunks = [text]
        
        # Build the quality maximization prompt
        def create_quality_maximization_prompt(chunk_text, chunk_num=None, total_chunks=None, has_assessment=False):
            chunk_info = f" (CHUNK {chunk_num} OF {total_chunks})" if chunk_num else ""
            
            assessment_context = ""
            if has_assessment and assessment_results:
                assessment_context = f"\n\nPREVIOUS QUALITY ASSESSMENT RESULTS:\n{assessment_results}\n\nUse these assessment results to guide your rewrite and address any weaknesses identified.\n"
            
            return f"""You are tasked with rewriting the following text{chunk_info} to MAXIMIZE its quality across all dimensions.

DEFAULT QUALITY MAXIMIZATION INSTRUCTIONS:
- Make sure ALL claims are clearly stated, adequately developed, and cogently argued for
- If claims are EMPIRICAL, cite the relevant empirical/experimental/historical information
- If claims are NON-EMPIRICAL (analytic), provide the relevant argumentation
- If the work already adequately states, develops, and substantiates its claims, then ensure the output gets maximally high scores on ALL quality dimensions below

QUALITY DIMENSIONS TO MAXIMIZE:

## STRUCTURAL AND STYLISTIC QUALITY
- Maintain coherent structure from start to finish
- Ensure smooth, logically motivated transitions between sections
- Make each paragraph contribute clearly to the central argument or theme
- Achieve appropriate pacing — neither rushed nor meandering
- Exhibit rhythm and stylistic control at the sentence level

## CLARITY AND READABILITY
- Clearly define and consistently use key concepts and terms
- Communicate complex ideas without unnecessary jargon
- Achieve optimal balance between concision and elaboration
- Ensure prose flows naturally, not stilted or mechanical
- Make argument easy to follow without external guidance

## SUBSTANTIVE AND EVIDENTIAL QUALITY
- Support claims with adequate reasoning, examples, or data
- Ensure evidence genuinely connects to the claims it supports
- Acknowledge or integrate counterarguments where appropriate
- Balance generalization and specificity appropriately
- Display awareness of limitations or assumptions

## VOICE AND ENGAGEMENT
- Convey confidence and authority without bluster
- Match tone to subject matter and intended audience
- Demonstrate stylistic distinctiveness, avoid generic or derivative prose
- Sustain reader engagement across full length
- Exhibit emotional or rhetorical intelligence (timing, emphasis, restraint)

## EMPIRICAL/HISTORICAL QUALITY
- Make factual claims verifiable and consistent with established data
- Distinguish clearly between evidence, interpretation, and speculation
- Accurately represent historical or experimental contexts (not retrofit to thesis)
- Derive argument naturally from data, not cherry-pick
- Support cause-effect relations rather than merely assert them
- Present statistical/quantitative reasoning correctly and transparently
- Acknowledge alternative explanations or competing evidence
- Maintain coherent, non-contradictory timelines and causal chains
- Show understanding of methodological limits or uncertainty
- Reveal genuine insight through evidence synthesis, not just fact accumulation

## LOGICAL/PHILOSOPHICAL/MATHEMATICAL QUALITY
- Define all key terms, variables, or premises precisely and use consistently
- Ensure inferences are valid (conclusions actually follow from premises)
- Maintain proper distinctions (necessary vs. sufficient, analytic vs. synthetic)
- Avoid equivocation and circular reasoning
- Anticipate and rigorously deal with counterexamples
- Ensure deductively or dialectically coherent progression (no leaps or gaps)
- Execute formal or quasi-formal structures accurately (proofs, schemata, analogies)
- Integrate examples or thought-experiments that genuinely test claims
- Connect theoretical claims to real conceptual or empirical stakes, not pure wordplay
- Demonstrate self-monitoring reasoning (aware of assumptions, scope, implications)

## INTELLIGENCE CRITERIA (from Intelligence Assessment)
- Make it profoundly insightful
- Develop all points fully and systematically
- Arrange ideas hierarchically with logical scaffolding
- Demonstrate masterful use of logic and reasoning
- Ensure all points are fresh and original, not clichés
- Use technical language for precision, not obfuscation
- Ensure organic development of ideas
- Open up inquiry, not shut it off
- Display genuine intelligence in content
- Be authentic, not phony
- Exhibit complex, coherent sentence-level logic
- Organize around strong governing concepts
- Demonstrate system-level control over ideas
- Present real, fresh points (not institutional propaganda)
- Be direct, not evasive
- Progress by logical entailment, not by authority
- Use sources to develop ideas, not mask lack of ideas
{assessment_context}
INSTRUCTIONS FOR REWRITING:
- Completely rewrite the text to maximize quality on ALL dimensions above
- Preserve core ideas but elevate them dramatically
- Make every sentence demonstrate sophisticated thought and clear argumentation
- Support all claims with appropriate evidence or reasoning
- Create hierarchical logical structure
- Be bold, original, and intellectually rigorous
- Maintain or exceed original length while adding depth and substance
- Write in a voice that is confident, clear, and authoritative

ORIGINAL TEXT TO MAXIMIZE:
{chunk_text}

NOW PROVIDE THE MAXIMALLY HIGH-QUALITY REWRITE:"""
        
        # Create streaming response with chunking support
        from flask import stream_with_context
        import time
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                has_assessment = bool(assessment_results.strip())
                
                for i, chunk in enumerate(chunks, 1):
                    if total_chunks > 1:
                        chunk_header = f"\n\n{'='*50}\nQUALITY-MAXIMIZED REWRITE - CHUNK {i} OF {total_chunks}\n{'='*50}\n\n"
                        yield f"data: {chunk_header}\n\n"
                    
                    prompt = create_quality_maximization_prompt(chunk, i if total_chunks > 1 else None, total_chunks if total_chunks > 1 else None, has_assessment)
                    
                    with client.messages.stream(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=6000,
                        system="You are an expert academic writer and editor capable of transforming text to maximize quality across all dimensions: structural, stylistic, substantive, evidential, logical, and rhetorical. You produce clear, rigorous, well-argued prose that supports all claims with appropriate evidence and reasoning.",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7
                    ) as stream:
                        for text_chunk in stream.text_stream:
                            # Send each chunk as SSE format
                            yield f"data: {text_chunk}\n\n"
                    
                    # Add delay between chunks to avoid rate limits
                    if i < total_chunks:
                        yield f"data: \n\n\n\n"
                        yield f"data: [Pausing 3 seconds to avoid rate limits...]\n\n"
                        time.sleep(3)  # 3 second delay between chunks
                        yield f"data: \n\n"
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in quality maximization: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/fiction_maximization_stream', methods=['POST'])
def fiction_maximization_stream():
    """
    Fiction Quality Maximization - Rewrites fiction to maximize scores on all 50 fiction assessment criteria with streaming response and chunking for large texts.
    """
    try:
        from anthropic import Anthropic
        import os
        import re
        
        data = request.get_json()
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Get Anthropic API key
        anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
        if not anthropic_key:
            return jsonify({'error': 'Anthropic API key not configured'}), 500
        
        # Initialize Anthropic client
        client = Anthropic(api_key=anthropic_key)
        
        # Chunk text if it's too large (split at ~2000 words per chunk to avoid token limits)
        words = text.split()
        word_count = len(words)
        chunks = []
        
        if word_count > 2000:
            # Split into chunks at paragraph boundaries
            paragraphs = re.split(r'\n\s*\n', text)
            current_chunk = []
            current_word_count = 0
            
            for para in paragraphs:
                para_words = len(para.split())
                if current_word_count + para_words > 2000 and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_word_count = para_words
                else:
                    current_chunk.append(para)
                    current_word_count += para_words
            
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
        else:
            chunks = [text]
        
        # Build the fiction maximization prompt
        def create_fiction_maximization_prompt(chunk_text, chunk_num=None, total_chunks=None):
            chunk_info = f" (CHUNK {chunk_num} OF {total_chunks})" if chunk_num else ""
            
            return f"""You are tasked with rewriting the following fiction text{chunk_info} to MAXIMIZE its quality across ALL fiction assessment criteria. Transform this into the best possible fiction that scores maximally high on all 50 dimensions below.

ORIGINAL FICTION TEXT:
{chunk_text}

---

FICTION QUALITY MAXIMIZATION CRITERIA - Score Maximally on ALL:

## A. NARRATIVE STRUCTURE AND COHERENCE
1. Create a clear structural arc (setup, conflict, resolution) - make the story architecture flawless
2. Ensure all plot events are causally connected, not random or episodic - every event must flow inevitably from what came before
3. Sustain perfect pacing with tension and momentum throughout
4. Make all scene transitions smooth and deeply meaningful
5. Achieve complete narrative unity - every scene must serve the central story or theme
6. Begin and end at exactly the right points - perfect story boundaries
7. Handle all exposition naturally, eliminating any info-dumps
8. Escalate stakes intelligently and organically, never through contrivance
9. If using time shifts (flashbacks, nonlinear structure), handle them with crystalline clarity
10. Maintain the story's point of view (first, third, omniscient, etc.) with absolute consistency

## B. CHARACTERIZATION AND PSYCHOLOGY
11. Make all characters psychologically believable and internally consistent - they must feel real
12. Reveal characters through vivid action and natural dialogue, never through exposition
13. Ensure motivations are crystal clear yet complex, never simplistic
14. Show main characters changing or deepening profoundly over the course of the story
15. Make all secondary characters distinct and purposeful, eliminate any filler
16. Build in rich moral or psychological ambiguity where appropriate
17. Develop character relationships with exquisite nuance and authentic tension
18. Ensure inner conflicts powerfully reflect larger thematic tensions
19. Craft speech patterns and thoughts perfectly suited to each character's background and situation
20. Make characters' choices drive the plot organically - never through convenience or coincidence

## C. STYLE AND LANGUAGE
21. Demonstrate masterful prose with perfect rhythm, rich texture, and precise tonal control
22. Create vivid, fresh images that are never overwrought or clichéd
23. Choose diction that fits the narrative voice and setting perfectly
24. Use figurative language (metaphor, simile, symbol) purposefully, never merely decoratively
25. Vary sentence structure expertly to reflect emotional shifts and control pacing
26. Write dialogue that sounds completely natural while revealing character depth
27. Achieve narrative precision - neither vague nor florid, always exact
28. Develop a distinctive, memorable voice or stylistic fingerprint
29. Maintain perfect internal tonal coherence (comic, tragic, lyrical, ironic, etc.)
30. Create language that evokes powerful sensory and emotional immediacy

## D. THEMATIC AND CONCEPTUAL DEPTH
31. Explore meaningful ideas and profound human truths
32. Integrate themes seamlessly into the narrative rather than imposing them from outside
33. Offer deep insight into psychology, society, or morality beyond the surface plot
34. Develop symbols or motifs consistently and with sophisticated intelligence
35. Resist all trivial moralizing or clichéd resolution
36. Create rich interpretive depth - enable multiple plausible readings
37. Make conflicts embody deeper philosophical or existential questions
38. Embody theme through character and action, never merely stating it
39. Invite genuine reflection without forcing it
40. Achieve satisfying closure without reducing the story's complexity

## E. EMOTIONAL AND AESTHETIC IMPACT
41. Evoke profound, authentic emotional engagement
42. Earn all emotional beats through deep character development and compelling situation
43. Achieve perfect balance between intellect and feeling
44. Create a masterful rhythm of tension and release that sustains rapt attention
45. Make the story unforgettable - create indelible scenes, lines, or insights
46. Make the story feel vibrantly alive - surprising yet inevitable
47. Leave the reader fundamentally altered, challenged, or deeply moved
48. Build a story world that is completely immersive and internally consistent
49. Create profound aesthetic pleasure beyond mere plot satisfaction
50. Achieve true resonance - a sense of meaning that transcends the literal

REWRITING INSTRUCTIONS:
- Transform this fiction to achieve MAXIMUM scores on ALL 50 criteria above
- Preserve the core story but elevate it to literary excellence
- Every sentence must demonstrate masterful craft
- Create flawless narrative architecture with deep thematic resonance
- Develop fully realized, psychologically complex characters
- Write prose that is both beautiful and precise
- Build emotional power that is completely earned
- Maintain or expand length while adding depth, richness, and artistry
- Write fiction that feels alive, inevitable, and unforgettable

NOW PROVIDE THE MAXIMALLY EXCELLENT FICTION REWRITE:"""
        
        # Create streaming response with chunking support
        from flask import stream_with_context
        import time
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                
                for i, chunk in enumerate(chunks, 1):
                    if total_chunks > 1:
                        chunk_header = f"\n\n{'='*50}\nFICTION QUALITY-MAXIMIZED REWRITE - CHUNK {i} OF {total_chunks}\n{'='*50}\n\n"
                        yield f"data: {chunk_header}\n\n"
                    
                    prompt = create_fiction_maximization_prompt(chunk, i if total_chunks > 1 else None, total_chunks if total_chunks > 1 else None)
                    
                    with client.messages.stream(
                        model="claude-sonnet-4-5-20250929",
                        max_tokens=6000,
                        system="You are a master fiction writer and editor with the ability to transform ordinary fiction into literary excellence. You craft narratives with flawless structure, psychologically rich characters, beautiful prose, deep thematic resonance, and profound emotional impact. Every element you write serves the story's artistic vision.",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.8
                    ) as stream:
                        for text_chunk in stream.text_stream:
                            # Send each chunk as SSE format
                            yield f"data: {text_chunk}\n\n"
                    
                    # Add delay between chunks to avoid rate limits
                    if i < total_chunks:
                        yield f"data: \n\n\n\n"
                        yield f"data: [Pausing 3 seconds to avoid rate limits...]\n\n"
                        time.sleep(3)  # 3 second delay between chunks
                        yield f"data: \n\n"
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in fiction quality maximization: {str(e)}")
        return jsonify({'error': str(e)}), 500


def clean_markdown(text):
    """
    Strip Markdown formatting so output renders cleanly in a plain textarea.
    Preserves LaTeX math spans (\\(...\\) and \\[...\\]).
    """
    import re

    # Protect LaTeX math so we don't mangle it
    placeholders = []

    def protect(match):
        placeholders.append(match.group(0))
        return f"\x00{len(placeholders) - 1}\x00"

    text = re.sub(r'\\\(.*?\\\)', protect, text, flags=re.DOTALL)
    text = re.sub(r'\\\[.*?\\\]', protect, text, flags=re.DOTALL)

    # Remove fenced code blocks markers and inline code backticks
    text = re.sub(r'```[a-zA-Z0-9]*\n?', '', text)
    text = text.replace('`', '')

    # Bold / italic emphasis
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'__(.*?)__', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\*(.*?)\*', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'(?<!\w)_(.*?)_(?!\w)', r'\1', text, flags=re.DOTALL)

    # Heading markers, bullet markers, blockquotes (line-based)
    text = re.sub(r'^\s{0,3}#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^(\s*)[*+]\s+', r'\1- ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s{0,3}>\s?', '', text, flags=re.MULTILINE)

    # Restore protected math spans
    for i, span in enumerate(placeholders):
        text = text.replace(f"\x00{i}\x00", span)

    return text


def force_paragraph_formatting(text):
    """
    FORCE paragraph breaks every 4 sentences - NO EXCEPTIONS.
    This function GUARANTEES proper formatting regardless of AI output.
    """
    import re
    
    # Split text into sentences (handles periods, question marks, exclamation points)
    sentences = re.split(r'([.!?]+(?:\s+|$))', text)
    
    # Reconstruct with sentences paired with their punctuation
    formatted_sentences = []
    for i in range(0, len(sentences)-1, 2):
        if sentences[i].strip():
            formatted_sentences.append(sentences[i] + sentences[i+1])
    
    # Group into paragraphs of MAX 4 sentences
    result = []
    current_paragraph = []
    
    for i, sentence in enumerate(formatted_sentences):
        current_paragraph.append(sentence.strip())
        
        # After every 4 sentences OR at the end, create a paragraph
        if (i + 1) % 4 == 0 or i == len(formatted_sentences) - 1:
            if current_paragraph:
                result.append(' '.join(current_paragraph))
                current_paragraph = []
    
    # Join paragraphs with double newlines (blank lines)
    return '\n\n'.join(result)


@app.route('/diagnostics')
def diagnostics_page():
    """One-button synthetic diagnostic dashboard."""
    return render_template('diagnostics.html')


_diagnostics_state = {'last_run': 0.0, 'lock': threading.Lock()}
DIAGNOSTICS_COOLDOWN_SECONDS = 15


@app.route('/run_diagnostics', methods=['POST'])
def run_diagnostics():
    """Run every API key + core function check and return structured results.

    Rate limited: each run makes live (paid) provider calls, so a short global
    cooldown prevents the endpoint from being hammered into cost amplification.
    """
    now = time.time()
    with _diagnostics_state['lock']:
        elapsed = now - _diagnostics_state['last_run']
        if elapsed < DIAGNOSTICS_COOLDOWN_SECONDS:
            wait = round(DIAGNOSTICS_COOLDOWN_SECONDS - elapsed, 1)
            return jsonify({'error': f'Please wait {wait}s before running diagnostics again.'}), 429
        _diagnostics_state['last_run'] = now

    try:
        from diagnostics import run_all_diagnostics
        return jsonify(run_all_diagnostics())
    except Exception as e:
        logger.error(f"Diagnostics run failed: {str(e)}")
        return jsonify({'error': f'Diagnostics failed: {str(e)}'}), 500


@app.route('/export_document', methods=['POST'])
def export_document():
    """
    Export the output text as a clean, downloadable file (txt, docx, or pdf).
    Markdown is stripped and paragraphs are preserved for professional formatting.
    """
    try:
        import html as html_module

        data = request.get_json(silent=True) or {}
        text = (data.get('text') or '').strip()
        fmt = (data.get('format') or 'txt').lower().strip()
        title = (data.get('title') or 'document').strip() or 'document'

        if not text:
            return jsonify({'error': 'No text to export'}), 400

        # Guard against runaway documents (~2M chars is far beyond any normal output)
        if len(text) > 2_000_000:
            return jsonify({'error': 'Text is too large to export'}), 400

        # Strip stray markdown so the downloaded file looks perfect
        text = clean_markdown(text)

        # Split into paragraphs on blank lines
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
        if not paragraphs:
            paragraphs = [text]

        safe_name = re.sub(r'[^A-Za-z0-9_-]+', '_', title).strip('_')[:60] or 'document'

        # ---- Plain text ----
        if fmt == 'txt':
            buffer = io.BytesIO()
            buffer.write(('\n\n'.join(paragraphs)).encode('utf-8'))
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype='text/plain; charset=utf-8',
                as_attachment=True,
                download_name=f'{safe_name}.txt'
            )

        # ---- Microsoft Word ----
        if fmt in ('docx', 'word', 'doc'):
            from docx import Document
            from docx.shared import Pt

            document = Document()
            normal = document.styles['Normal']
            normal.font.name = 'Calibri'
            normal.font.size = Pt(12)

            for para in paragraphs:
                p = document.add_paragraph()
                p.paragraph_format.space_after = Pt(10)
                p.paragraph_format.line_spacing = 1.5
                lines = para.split('\n')
                for idx, line in enumerate(lines):
                    if idx > 0:
                        p.add_run().add_break()
                    p.add_run(line)

            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'{safe_name}.docx'
            )

        # ---- PDF ----
        if fmt == 'pdf':
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_JUSTIFY

            buffer = io.BytesIO()
            pdf = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                leftMargin=0.9 * inch,
                rightMargin=0.9 * inch,
                topMargin=1 * inch,
                bottomMargin=1 * inch
            )
            styles = getSampleStyleSheet()
            body_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=11.5,
                leading=17,
                spaceAfter=12,
                alignment=TA_JUSTIFY
            )

            story = []
            for para in paragraphs:
                escaped = html_module.escape(para).replace('\n', '<br/>')
                story.append(Paragraph(escaped, body_style))
                story.append(Spacer(1, 4))

            pdf.build(story)
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f'{safe_name}.pdf'
            )

        return jsonify({'error': f'Unsupported format: {fmt}'}), 400

    except Exception as e:
        logger.error(f"Export document error: {str(e)}")
        return jsonify({'error': f'Export failed: {str(e)}'}), 500


@app.route('/customized_rewrite_stream', methods=['POST'])
def customized_rewrite_stream():
    """
    Customized Rewrite - Rewrites text based on user's custom instructions, or uses smart defaults for fiction vs non-fiction.
    """
    try:
        import os
        import re
        from ai_failover import generate_with_failover
        
        data = request.get_json()
        text = data.get('text', '')
        custom_instructions = data.get('custom_instructions', '').strip()
        author_style_slug = data.get('author_style_slug', '').strip()
        preferred_provider = data.get('preferred_provider', 'auto').strip()

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        # If caller passed an author style slug, look up its rich style_prompt from the DB
        # and prepend it to any custom instructions the user also wrote.
        if author_style_slug:
            try:
                from models import AuthorStyle
                style_obj = AuthorStyle.query.filter_by(slug=author_style_slug, is_active=True).first()
                if style_obj:
                    db_prompt = style_obj.style_prompt
                    if custom_instructions:
                        custom_instructions = db_prompt + "\n\nADDITIONAL INSTRUCTIONS:\n" + custom_instructions
                    else:
                        custom_instructions = db_prompt
                    logger.info(f"Applied author style '{style_obj.name}' from DB")
            except Exception as _ae:
                logger.warning(f"Could not load author style '{author_style_slug}': {_ae}")

        # Default instructions if none provided
        if not custom_instructions:
            custom_instructions = """If the text is fiction (narrative with characters, dialogue, scenes, plot):
- Maximize quality by developing the story with richer characterization, deeper thematic resonance, more vivid prose, stronger narrative structure, and enhanced emotional impact
- Create psychological depth, compelling dialogue, immersive scenes, and literary excellence

If the text is non-fiction (expository, argumentative, analytical):
- Maximize quality by delineating and substantiating points with rigorous logical reasoning and empirical support
- Strengthen arguments with evidence, clarify structure, enhance analytical depth, and provide thorough intellectual justification for all claims"""
        
        # Chunk text if it's too large (split at ~2000 words per chunk to avoid token limits)
        words = text.split()
        word_count = len(words)
        chunks = []
        
        if word_count > 2000:
            # Split into chunks at paragraph boundaries
            paragraphs = re.split(r'\n\s*\n', text)
            current_chunk = []
            current_word_count = 0
            
            for para in paragraphs:
                para_words = len(para.split())
                if current_word_count + para_words > 2000 and current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [para]
                    current_word_count = para_words
                else:
                    current_chunk.append(para)
                    current_word_count += para_words
            
            if current_chunk:
                chunks.append('\n\n'.join(current_chunk))
        else:
            chunks = [text]
        
        # Build the customized rewrite prompt
        def create_customized_rewrite_prompt(chunk_text, instructions, chunk_num=None, total_chunks=None):
            chunk_info = f" (CHUNK {chunk_num} OF {total_chunks})" if chunk_num else ""
            
            return f"""You are tasked with rewriting the following text{chunk_info} according to the specific instructions provided.

ORIGINAL TEXT:
{chunk_text}

---

REWRITE INSTRUCTIONS:
{instructions}

---

IMPORTANT GUIDELINES:
- Follow the instructions precisely
- Preserve all factual content and core ideas from the original
- Maintain or enhance the original length unless compression is specifically requested
- Ensure the rewrite is polished, coherent, and of high quality
- Begin your response directly with the rewritten text - NO preambles, meta-commentary, or explanations

CRITICAL FORMATTING REQUIREMENT - ABSOLUTELY MANDATORY - NO EXCEPTIONS:

YOU MUST NEVER WRITE MORE THAN 4 SENTENCES IN A ROW WITHOUT A BLANK LINE.

After EVERY 3-4 sentences, you MUST insert a blank line (two newline characters: \n\n).

MAXIMUM PARAGRAPH LENGTH: 4 SENTENCES. NO EXCEPTIONS WHATSOEVER.

Your output MUST look like this:

First sentence here. Second sentence. Third sentence. Fourth sentence.

New paragraph starts after blank line. Another sentence. Third sentence here. Fourth sentence ends it.

Next paragraph with max four sentences. Second sentence. Third one. Fourth concludes.

ABSOLUTE RULES:
- NEVER write 5 or more sentences without a blank line
- Each paragraph: 3-4 sentences MAXIMUM
- Always use blank lines between paragraphs
- NO giant blocks of text under any circumstances

NOW PROVIDE THE REWRITTEN TEXT WITH MANDATORY PARAGRAPH BREAKS EVERY 3-4 SENTENCES:"""
        
        # Create streaming response with chunking support
        from flask import stream_with_context
        import time
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                
                for i, chunk in enumerate(chunks, 1):
                    if total_chunks > 1:
                        chunk_header = f"\n\n{'='*50}\nCUSTOMIZED REWRITE - CHUNK {i} OF {total_chunks}\n{'='*50}\n\n"
                        yield f"data: {chunk_header}\n\n"
                    
                    prompt = create_customized_rewrite_prompt(chunk, custom_instructions, i if total_chunks > 1 else None, total_chunks if total_chunks > 1 else None)
                    
                    # Generate with automatic key + provider failover (Anthropic first for rewrites).
                    # Build provider order: put user's choice first, fallback to rest
                    _all = ['anthropic', 'openai', 'deepseek', 'azure', 'venice', 'perplexity']
                    if preferred_provider and preferred_provider != 'auto' and preferred_provider in _all:
                        _order = [preferred_provider] + [p for p in _all if p != preferred_provider]
                    else:
                        _order = _all

                    result = generate_with_failover(
                        prompt,
                        system="You are an expert writing editor and rewriter. You transform text precisely according to user instructions while maintaining quality, clarity, and coherence. You never add preambles or meta-commentary - you provide the rewritten text directly.",
                        max_tokens=6000,
                        temperature=0.7,
                        preferred_order=_order,
                    )
                    full_response = result['text']
                    
                    # NOW APPLY FORCED PARAGRAPH FORMATTING - GUARANTEED BREAKS EVERY 4 SENTENCES
                    formatted_response = force_paragraph_formatting(full_response)
                    
                    # Stream in chunks of ~50 characters to preserve ALL formatting including \n\n
                    chunk_size = 50
                    total_words = 0
                    
                    for i in range(0, len(formatted_response), chunk_size):
                        chunk = formatted_response[i:i+chunk_size]
                        yield f"data: {json.dumps(chunk)}\n\n"
                    
                    # Separator between chunks
                    if i < total_chunks:
                        yield f"data: \n\n\n\n"
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in customized rewrite: {str(e)}")
        return jsonify({'error': str(e)}), 500


# ===== AUTHOR STYLES API =====

@app.route('/api/author_styles', methods=['GET'])
def list_author_styles():
    """Return all active author styles for the style picker UI."""
    try:
        from models import AuthorStyle
        styles = AuthorStyle.query.filter_by(is_active=True).order_by(AuthorStyle.name).all()
        return jsonify([s.to_dict() for s in styles])
    except Exception as e:
        logger.error(f"Error listing author styles: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/author_styles/<slug>', methods=['GET'])
def get_author_style(slug):
    """Return a single author style by slug."""
    try:
        from models import AuthorStyle
        style = AuthorStyle.query.filter_by(slug=slug, is_active=True).first()
        if not style:
            return jsonify({'error': 'Author style not found'}), 404
        return jsonify(style.to_dict())
    except Exception as e:
        logger.error(f"Error fetching author style {slug}: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/author_styles', methods=['POST'])
def create_author_style():
    """Admin: add a new author style."""
    try:
        from models import AuthorStyle
        data = request.get_json()
        required = ['name', 'slug', 'style_prompt']
        for f in required:
            if not data.get(f):
                return jsonify({'error': f'Missing required field: {f}'}), 400
        style = AuthorStyle(
            name=data['name'],
            slug=data['slug'],
            genre=data.get('genre'),
            era=data.get('era'),
            description=data.get('description'),
            style_prompt=data['style_prompt'],
            sample_text=data.get('sample_text'),
            is_active=data.get('is_active', True),
        )
        db.session.add(style)
        db.session.commit()
        return jsonify(style.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating author style: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/author_styles/<slug>', methods=['PUT'])
def update_author_style(slug):
    """Admin: update an existing author style."""
    try:
        from models import AuthorStyle
        style = AuthorStyle.query.filter_by(slug=slug).first()
        if not style:
            return jsonify({'error': 'Author style not found'}), 404
        data = request.get_json()
        for field in ['name', 'genre', 'era', 'description', 'style_prompt', 'sample_text', 'is_active']:
            if field in data:
                setattr(style, field, data[field])
        db.session.commit()
        return jsonify(style.to_dict())
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating author style {slug}: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/author_styles/<slug>', methods=['DELETE'])
def delete_author_style(slug):
    """Admin: delete an author style."""
    try:
        from models import AuthorStyle
        style = AuthorStyle.query.filter_by(slug=slug).first()
        if not style:
            return jsonify({'error': 'Author style not found'}), 404
        db.session.delete(style)
        db.session.commit()
        return jsonify({'deleted': slug})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting author style {slug}: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/author_styles')
def author_styles_page():
    """Author styles management page."""
    return render_template('author_styles.html')


# ===== HUMANIZER ROUTES =====

# PRESET_TEXT mapping for instruction presets
PRESET_TEXT = {
    "Compression — light (−15%)": "Cut filler; merge short clauses; keep meaning. Target ≈15% shorter.",
    "Compression — medium (−30%)": "Trim hard; delete throat-clearing; tighten syntax. Target ≈30% shorter.",
    "Compression — heavy (−45%)": "Sever redundancies; collapse repeats; keep core claims. Target ≈45% shorter.",
    "Mixed cadence": "Alternate short (5–12 words) and long (20–35 words) sentences; avoid uniform rhythm.",
    "Clause surgery": "Reorder main/subordinate clauses in ~30% of sentences without changing meaning.",
    "Front-load claim": "Put the main conclusion in sentence 1; evidence follows.",
    "Back-load claim": "Delay the main conclusion to the final 2–3 sentences.",
    "Seam/pivot": "Drop smooth connectors once; allow one abrupt thematic pivot.",
    "Imply one step": "Omit one obvious inferential step; keep it implicit (context makes it recoverable).",
    "Conditional framing": "Recast one key sentence as: If/Unless …, then …. Keep content identical.",
    "Local contrast": "Use exactly one contrast marker (but/except/aside) to mark a boundary; add no new facts.",
    "Scope check": "Replace one absolute with a bounded form (e.g., 'in cases like these').",
    "Deflate jargon": "Swap nominalizations for plain verbs where safe (e.g., utilization→use).",
    "Kill stock transitions": "Delete 'Moreover/Furthermore/In conclusion' everywhere.",
    "Hedge once": "Use exactly one hedge: probably/roughly/more or less.",
    "Drop intensifiers": "Remove 'very/clearly/obviously/significantly'.",
    "Low-heat voice": "Prefer plain verbs; avoid showy synonyms.",
    "One aside": "Allow one short parenthetical or em-dash aside; strictly factual.",
    "Concrete benchmark": "Replace one vague scale with a testable one (e.g., 'enough to X').",
    "Swap generic example": "If the source has an example, make it slightly more specific; else skip.",
    "Metric nudge": "Replace 'more/better' with a minimal, source-safe comparator (e.g., 'more than last case').",
    "Asymmetric emphasis": "Linger on the main claim; compress secondary points sharply.",
    "Cull repeats": "Delete duplicated sentences/ideas; keep the strongest instance.",
    "Topic snap": "Allow one abrupt focus change; no recap.",
    "No lists": "Output as continuous prose; remove bullets/numbering.",
    "No meta": "No prefaces/apologies/phrases like 'as requested'.",
    "Exact nouns": "Replace ambiguous pronouns with exact nouns.",
    "Quote once": "If the source has a strong phrase, quote it once; otherwise skip.",
    "Claim lock": "Do not add examples, scenarios, or data not present in the source.",
    "Entity lock": "Keep names, counts, and attributions exactly as given.",
    # Combo presets expand to atomic ones:
    "Lean & Sharp": "Compression — medium (−30%); Mixed cadence; Imply one step; Kill stock transitions",
    "Analytic": "Clause surgery; Front-load claim; Scope check; Exact nouns; No lists",
}

# Default style sample: FORMAL AND FUNCTIONAL RELATIONSHIPS (Content-Neutral category)
FORMAL_FUNCTIONAL = """There are two broad types of relationships: formal and functional.
Formal relationships hold between descriptions. A description is any statement that can be true or false.
Example of a formal relationship: The description that a shape is a square cannot be true unless the description that it has four equal sides is true. Therefore, a shape's being a square depends on its having four equal sides.

Functional relationships hold between events or conditions. (An event is anything that happens in time.)
Example of a functional relationship: A plant cannot grow without water. Therefore, a plant's growth depends on its receiving water.

The first type is structural, i.e., it holds between statements about features.
The second is operational, i.e., it holds between things in the world as they act or change.

Descriptions as objects of consideration
The objects of evaluation are descriptions. Something is not evaluated unless it is described, and it is not described unless it can be stated. One can notice non-descriptions — sounds, objects, movements — but in the relevant sense one evaluates descriptions of them.

Relationships not known through direct observation
Some relationships are known, not through direct observation, but through reasoning. Such relationships are structural, as opposed to observational. Examples of structural relationships are:

If A, then A or B.

All tools require some form of use."""

# Legacy: Raven Paradox style sample (kept for reference)
RAVEN_PARADOX = """Presumably, logically equivalent statements are confirmationally equivalent. In other words, if two statements entail each other, then anything that one confirms the one statement to a given degree also confirms the other statement to that degree. But this actually seems false when consider statement-pairs such as: 

(i) All ravens are black, 
and 
(ii) All non-black things are non-ravens, 

which, though logically equivalent, seem to confirmationally equivalent, in that a non-black non-raven confirms (ii) to a high degree but confirms (i) to no degree or at most to a low degree. 
A number of very contrived solutions to this paradox have been proposed, all of which either deny that there is a paradox or invent ad hoc systems of logic to validate the 'solution' in question. 
But the real solution is clear. First of all, it is only principled generalizations that can be confirmed. Supposing that you assert (i) with the intention of affirming a principled as opposed to an accidental generalization, you are saying that instances of the property of being a raven grounds or causes instances of blackness. Read thus, (i) is most certainly not equivalent with (ii) or with any variation thereof. Be it noted that while there is a natural nomic or causal reading of (i), there is no such reading of (ii). Also be it noted that it is only principled as opposed to accidental generalizations that can be confirmed. "All metal expands when heated" can be confirmed but not "all objects in Smith's pocket expand when heated." In general, when read as principled and therefore confirmable generalization, "all x's are y's" has nomic or causal content is therefore not equivalent with "all non-y's are non-x's." Case closed on the Raven Paradox."""

def expand_presets(selected):
    """Expand combo presets into atomic presets"""
    out = []
    seen = set()
    
    for name in selected:
        txt = PRESET_TEXT.get(name)
        if not txt:
            continue
            
        # Check if combo preset (contains ; but not …)
        if ';' in txt and '…' not in txt:
            # Split by ; and add atomic presets
            parts = [s.strip() for s in txt.split(';')]
            for alias in parts:
                if alias in PRESET_TEXT and alias not in seen:
                    seen.add(alias)
                    out.append(alias)
        else:
            if name not in seen:
                seen.add(name)
                out.append(name)
    
    return out

def build_preset_block(selected_presets, custom_instructions):
    """Build the preset instructions block"""
    expanded = expand_presets(selected_presets or [])
    lines = []
    
    for name in expanded:
        lines.append(f"- {PRESET_TEXT[name]}")
    
    if custom_instructions and custom_instructions.strip():
        lines.append(f"- {custom_instructions.strip()}")
    
    if not lines:
        return ""
    
    return f"Apply ONLY these additional rewrite instructions (no other goals):\n" + "\n".join(lines) + "\n\n"

def build_rewrite_prompt(input_text, style_text, selected_presets, custom_instructions):
    """Build the complete rewrite prompt for 3-box Humanizer with SURGICAL PRECISION style cloning"""
    has_style = bool(style_text and style_text.strip())
    
    # Use provided style or default to FORMAL AND FUNCTIONAL RELATIONSHIPS
    style_sample = style_text.strip() if has_style else FORMAL_FUNCTIONAL
    
    # Build surgical precision style-cloning prompt
    prompt = f'''STYLE SAMPLE TO CLONE (analyze at molecular level):
{style_sample}

CRITICAL INSTRUCTIONS FOR SURGICAL PRECISION STYLE CLONING:

1. ANALYZE the style sample's specific characteristics:
   - Sentence length patterns (short vs long, variation)
   - Clause structure (how ideas are connected, subordination patterns)
   - Word choice (formal vs informal, technical vs plain, concrete vs abstract)
   - Rhythm and cadence (how sentences flow)
   - Punctuation patterns (em-dashes, semicolons, commas)
   - Paragraph structure
   - Transitional patterns
   - Level of elaboration vs compression

2. PRESERVE the meaning and content of the input text EXACTLY - do not add new facts or examples

3. REWRITE the input text by applying the EXACT stylistic patterns from the style sample:
   - Match sentence length distribution
   - Copy clause arrangement patterns
   - Use similar vocabulary register
   - Mirror the rhythm and flow
   - Apply the same punctuation style
   - Maintain similar elaboration density

4. OUTPUT REQUIREMENTS:
   - Provide ONLY the rewritten text - NO preambles, NO explanations
   - Do NOT say "Here's the rewrite" or similar
   - Start directly with the rewritten content
   - The rewrite must pass AI detection as human-written
'''
    
    # Add preset instructions
    preset_block = build_preset_block(selected_presets, custom_instructions)
    if preset_block:
        prompt += "\n" + preset_block
    
    # Add input text
    prompt += f'\nINPUT TEXT TO REWRITE:\n{input_text}\n\nREWRITTEN TEXT:'
    
    return prompt

@app.route('/humanizer_rewrite_stream', methods=['POST'])
def humanizer_rewrite_stream():
    """Streaming rewrite endpoint for Humanizer with 3-box system (Input, Style, Output)"""
    try:
        data = request.json
        input_text = data.get('inputText', '').strip()
        style_text = data.get('styleText', '').strip()
        selected_presets = data.get('selectedPresets', [])
        custom_instructions = data.get('customInstructions', '').strip()
        provider = data.get('provider', 'anthropic')
        
        if not input_text:
            return jsonify({'error': 'No input text provided'}), 400
        
        # Build the rewrite prompt
        prompt = build_rewrite_prompt(input_text, style_text, selected_presets, custom_instructions)
        
        # Chunk text for large inputs (2000 words with 100 word overlap)
        words = input_text.split()
        word_count = len(words)
        chunks = []
        chunk_size = 2000
        overlap = 100
        
        if word_count > chunk_size:
            # Create overlapping chunks
            i = 0
            while i < word_count:
                end = min(i + chunk_size, word_count)
                chunk = ' '.join(words[i:end])
                chunks.append(chunk)
                i += chunk_size - overlap if i + chunk_size < word_count else word_count
        else:
            chunks = [input_text]
        
        # Create streaming response
        from flask import stream_with_context
        
        @stream_with_context
        def generate():
            try:
                total_chunks = len(chunks)
                
                # Provider-specific implementations
                if provider == 'anthropic':
                    anthropic_key = os.environ.get('ANTHROPIC_API_KEY')
                    if not anthropic_key:
                        yield f"data: Error: Anthropic API key not configured\n\n"
                        return
                    
                    client = Anthropic(api_key=anthropic_key)
                    
                    for i, chunk in enumerate(chunks, 1):
                        # Build prompt for this chunk
                        chunk_prompt = build_rewrite_prompt(chunk, style_text, selected_presets, custom_instructions)
                        
                        # Progress update
                        progress = int((i - 1) / total_chunks * 100)
                        yield f"data: PROGRESS:{progress}\n\n"
                        
                        if total_chunks > 1:
                            yield f"data: \n\n[CHUNK {i}/{total_chunks}]\n\n"
                        
                        # Stream from Anthropic
                        with client.messages.stream(
                            model="claude-sonnet-4-5-20250929",
                            max_tokens=6000,
                            system="You are a precision text rewriter. Output ONLY the rewritten text with NO preambles, NO explanations, NO meta-commentary. Start directly with the rewritten content.",
                            messages=[{"role": "user", "content": chunk_prompt}],
                            temperature=0.7
                        ) as stream:
                            for text_chunk in stream.text_stream:
                                yield f"data: {text_chunk}\n\n"
                        
                        # Delay between chunks
                        if i < total_chunks:
                            yield f"data: \n\n\n\n"
                            time.sleep(3)
                
                elif provider == 'openai':
                    openai_key = os.environ.get('OPENAI_API_KEY')
                    if not openai_key:
                        yield f"data: Error: OpenAI API key not configured\n\n"
                        return
                    
                    from openai import OpenAI
                    client = OpenAI(api_key=openai_key)
                    
                    for i, chunk in enumerate(chunks, 1):
                        chunk_prompt = build_rewrite_prompt(chunk, style_text, selected_presets, custom_instructions)
                        progress = int((i - 1) / total_chunks * 100)
                        yield f"data: PROGRESS:{progress}\n\n"
                        
                        if total_chunks > 1:
                            yield f"data: \n\n[CHUNK {i}/{total_chunks}]\n\n"
                        
                        stream = client.chat.completions.create(
                            model="gpt-4",
                            messages=[{"role": "user", "content": chunk_prompt}],
                            stream=True,
                            max_tokens=6000,
                            temperature=0.7
                        )
                        
                        for completion in stream:
                            if completion.choices[0].delta.content:
                                yield f"data: {completion.choices[0].delta.content}\n\n"
                        
                        if i < total_chunks:
                            yield f"data: \n\n\n\n"
                            time.sleep(3)
                
                elif provider == 'deepseek':
                    deepseek_key = os.environ.get('DEEPSEEK_API_KEY')
                    if not deepseek_key:
                        yield f"data: Error: DeepSeek API key not configured\n\n"
                        return
                    
                    from openai import OpenAI
                    client = OpenAI(api_key=deepseek_key, base_url="https://api.deepseek.com")
                    
                    for i, chunk in enumerate(chunks, 1):
                        chunk_prompt = build_rewrite_prompt(chunk, style_text, selected_presets, custom_instructions)
                        progress = int((i - 1) / total_chunks * 100)
                        yield f"data: PROGRESS:{progress}\n\n"
                        
                        if total_chunks > 1:
                            yield f"data: \n\n[CHUNK {i}/{total_chunks}]\n\n"
                        
                        stream = client.chat.completions.create(
                            model="deepseek-chat",
                            messages=[{"role": "user", "content": chunk_prompt}],
                            stream=True,
                            max_tokens=6000,
                            temperature=0.7
                        )
                        
                        for completion in stream:
                            if completion.choices[0].delta.content:
                                yield f"data: {completion.choices[0].delta.content}\n\n"
                        
                        if i < total_chunks:
                            yield f"data: \n\n\n\n"
                            time.sleep(3)
                
                elif provider == 'perplexity':
                    perplexity_key = os.environ.get('PERPLEXITY_API_KEY')
                    if not perplexity_key:
                        yield f"data: Error: Perplexity API key not configured\n\n"
                        return
                    
                    from openai import OpenAI
                    client = OpenAI(api_key=perplexity_key, base_url="https://api.perplexity.ai")
                    
                    for i, chunk in enumerate(chunks, 1):
                        chunk_prompt = build_rewrite_prompt(chunk, style_text, selected_presets, custom_instructions)
                        progress = int((i - 1) / total_chunks * 100)
                        yield f"data: PROGRESS:{progress}\n\n"
                        
                        if total_chunks > 1:
                            yield f"data: \n\n[CHUNK {i}/{total_chunks}]\n\n"
                        
                        stream = client.chat.completions.create(
                            model="sonar",
                            messages=[{"role": "user", "content": chunk_prompt}],
                            stream=True,
                            max_tokens=6000,
                            temperature=0.7
                        )
                        
                        for completion in stream:
                            if completion.choices[0].delta.content:
                                yield f"data: {completion.choices[0].delta.content}\n\n"
                        
                        if i < total_chunks:
                            yield f"data: \n\n\n\n"
                            time.sleep(3)
                
                # Final progress
                yield f"data: PROGRESS:100\n\n"
                
            except Exception as e:
                logger.error(f"Humanizer streaming error: {str(e)}")
                yield f"data: \n\nError: {str(e)}\n\n"
        
        return Response(generate(), mimetype='text/event-stream', headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        })
        
    except Exception as e:
        logger.error(f"Error in humanizer rewrite: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_humanizer_docx', methods=['POST'])
def download_humanizer_docx():
    """Download humanizer output as DOCX"""
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Create DOCX document
        doc = Document()
        doc.add_heading('Humanizer Output', 0)
        
        # Add paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to BytesIO
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        return send_file(
            docx_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'humanizer_output_{int(time.time())}.docx'
        )
        
    except Exception as e:
        logger.error(f"Error downloading DOCX: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_humanizer_pdf', methods=['POST'])
def download_humanizer_pdf():
    """Download humanizer output as PDF"""
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Create PDF
        pdf_io = io.BytesIO()
        c = canvas.Canvas(pdf_io)
        
        # Add title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 800, "Humanizer Output")
        
        # Add text with word wrapping
        c.setFont("Helvetica", 11)
        y = 760
        max_width = 500
        
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                words = para.strip().split()
                line = ""
                for word in words:
                    test_line = line + word + " "
                    if c.stringWidth(test_line, "Helvetica", 11) < max_width:
                        line = test_line
                    else:
                        c.drawString(50, y, line.strip())
                        y -= 15
                        line = word + " "
                        if y < 50:
                            c.showPage()
                            c.setFont("Helvetica", 11)
                            y = 800
                if line:
                    c.drawString(50, y, line.strip())
                    y -= 20
                    if y < 50:
                        c.showPage()
                        c.setFont("Helvetica", 11)
                        y = 800
        
        c.save()
        pdf_io.seek(0)
        
        return send_file(
            pdf_io,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'humanizer_output_{int(time.time())}.pdf'
        )
        
    except Exception as e:
        logger.error(f"Error downloading PDF: {str(e)}")
        return jsonify({'error': str(e)}), 500

